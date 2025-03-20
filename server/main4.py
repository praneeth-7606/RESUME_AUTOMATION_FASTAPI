# from fastapi import FastAPI, File, UploadFile, HTTPException, Form
# from pydantic import BaseModel
# import fitz  # PyMuPDF for PDF files
# from docx import Document  # For DOCX files
# import pandas as pd
# import json
# import shutil
# from mistralai import Mistral
# import os

# app = FastAPI()

# # Mistral API Key
# API_KEY = "QsIe8weZaSTfNx9ykWlbjC3vGhWVSzWE"
# client = Mistral(api_key=API_KEY)

# # Function to extract text from uploaded resume
# def extract_text(file_path: str) -> str:
#     ext = file_path.split(".")[-1].lower()
#     text = ""
#     if ext == "pdf":
#         doc = fitz.open(file_path)
#         for page in doc:
#             text += page.get_text("text")
#     elif ext == "docx":
#         doc = Document(file_path)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#     elif ext == "txt":
#         with open(file_path, "r", encoding="utf-8") as file:
#             text = file.read()
#     else:
#         raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF, DOCX, or TXT.")
#     return text

# # Function to process skill matrix Excel file
# def process_skill_matrix(file_path: str):
#     xls = pd.ExcelFile(file_path)
#     sheets_data = []
#     record_id = 1
#     for sheet_name in xls.sheet_names[1:]:
#         df = pd.read_excel(xls, sheet_name=sheet_name)
#         df.columns = ['First Name', 'Last Name', 'Experience', 'Expertise'] + list(df.columns[4:])
#         records = []
#         for _, row in df.iterrows():
#             record = {"ID": record_id, **row.to_dict()}
#             records.append(record)
#             record_id += 1
#         sheets_data.append({"Sheet Name": sheet_name, "Data": records})
#     return json.dumps(sheets_data, indent=4)

# # FastAPI endpoint to upload resume and extract text
# @app.post("/upload_resume/")
# async def upload_resume(file: UploadFile = File(...)):
#     file_location = f"temp_{file.filename}"
#     with open(file_location, "wb") as buffer:
#         shutil.copyfileobj(file.file, buffer)
#     extracted_text = extract_text(file_location)
#     os.remove(file_location)
#     return {"extracted_text": extracted_text}

# # FastAPI endpoint to upload skill matrix and process data
# @app.post("/upload_skill_matrix/")
# async def upload_skill_matrix(file: UploadFile = File(...)):
#     file_location = f"temp_{file.filename}"
#     with open(file_location, "wb") as buffer:
#         shutil.copyfileobj(file.file, buffer)
#     skill_matrix_json = process_skill_matrix(file_location)
#     os.remove(file_location)
#     return json.loads(skill_matrix_json)

# # FastAPI endpoint to generate structured resume JSON
# @app.post("/generate_resume/")
# async def generate_resume(extracted_text: str = Form(...), skill_matrix: str = Form(...)):
#     system_prompt = "You are an AI assistant that reformats resumes into structured JSON format."
#     response = client.chat.complete(
#         model="mistral-large-latest",
#         messages=[
#             {"role": "system", "content": system_prompt},
#             {"role": "user", "content": f"Extracted Resume Text: {extracted_text}\nSkill Matrix: {skill_matrix}\nFormat into JSON"}
#         ],
#         response_format={"type": "json_object"}
#     )
#     return json.loads(response.choices[0].message.content)

# # FastAPI endpoint to generate cover letter
# @app.post("/generate_cover_letter/")
# async def generate_cover_letter(resume_summary: str = Form(...)):
#     system_prompt = "You are an AI assistant that generates professional cover letters."
#     response = client.chat.complete(
#         model="mistral-large-latest",
#         messages=[
#             {"role": "system", "content": system_prompt},
#             {"role": "user", "content": f"Resume Summary: {resume_summary}\nGenerate a professional cover letter."}
#         ],
#         response_format={"type": "text"}
#     )
#     return {"cover_letter": response.choices[0].message.content}

# # Run the FastAPI app with uvicorn (for local testing)
# # Command: uvicorn filename:app --reload
# from fastapi import FastAPI, File, UploadFile, Form, HTTPException, BackgroundTasks
# from fastapi.responses import JSONResponse, FileResponse
# from typing import List, Optional, Dict, Any
# import pandas as pd
# import json
# import fitz  # PyMuPDF
# from docx import Document
# import io
# import os
# from pydantic import BaseModel
# from mistralai import Mistral
# from fpdf import FPDF

# app = FastAPI(title="Resume Automation API")

# # API key for Mistral
# API_KEY = "QsIe8weZaSTfNx9ykWlbjC3vGhWVSzWE"

# # Global variable to store loaded skill matrix data
# sheets_data = []

# # Paths configuration
# BASE_PATH = os.path.dirname(os.path.abspath(__file__))
# LOGO_PATH = os.path.join(BASE_PATH, "Logo.png")
# PROMPTS_DIR = os.path.join(BASE_PATH, "Prompts")
# OUTPUT_DIR = os.path.join(BASE_PATH, "Output")

# # Create output directory if it doesn't exist
# os.makedirs(OUTPUT_DIR, exist_ok=True)

# # Pydantic models
# class NameQuery(BaseModel):
#     first_name: str
#     last_name: str

# class ResumeRequest(BaseModel):
#     old_resume_path: str
#     template_path: str
#     skill_matrix_path: Optional[str] = None
#     first_name: Optional[str] = None
#     last_name: Optional[str] = None

# class PDF(FPDF):
#     def header(self):
#         """Adds a logo and title at the top of the page."""
#         self.image(LOGO_PATH, 160, 8, 35)
#         self.set_font("Arial", "B", 12)
#         self.ln(10)

#     def chapter_title(self, title):
#         """Formats section titles with blue color."""
#         self.set_text_color(0, 0, 180)  # Set heading color to blue
#         self.set_font("Arial", "B", 11)
#         self.cell(0, 8, title, ln=True, align="L")
#         self.ln(4)
#         self.set_text_color(0, 0, 0)  # Reset color to black for body text
            
#     def add_project_table(self, project):
#         """Creates a vertical-style project details table with aligned description box"""
#         col_widths = [50, 120]  # Adjust column widths
#         row_height = 8  # Default row height

#         # Project Title (Bold)
#         self.set_font("Arial", "B", 9)
#         self.cell(0, 6, f"<{project['name']}>", ln=True)
#         self.ln(2)

#         # Set normal font for table content
#         self.set_font("Arial", "", 9)

#         # Project Name
#         self.cell(col_widths[0], row_height, "Project Name", border=1, align="C")
#         self.cell(col_widths[1], row_height, project["name"], border=1, ln=True)

#         # Role
#         self.cell(col_widths[0], row_height, "Role", border=1, align="C")
#         self.cell(col_widths[1], row_height, project["role"], border=1, ln=True)

#         # Calculate the height of the Description text
#         description_text = project["description"]
#         description_height = self.get_string_width(description_text) // col_widths[1] * row_height + row_height

#         # Description Title (Ensures it matches the height of the text)
#         self.cell(col_widths[0], description_height, "Description", border=1, align="C")
#         self.multi_cell(col_widths[1], row_height, description_text, border=1)

#         # Technology Row
#         self.cell(col_widths[0], row_height, "Technology", border=1, align="C")
#         self.multi_cell(col_widths[1], row_height, project["technology"], border=1)

#         role_played_text = project["role_played"]
#         role_played_height = self.get_string_width(role_played_text) // col_widths[1] * row_height + row_height

#         # Role Played
#         self.cell(col_widths[0], role_played_height, "Role Played", border=1, align="C")
#         self.multi_cell(col_widths[1], row_height, role_played_text, border=1)

#         self.ln()

# # File extraction functions
# def extract_text_from_file(file_path: str) -> str:
#     """Extract text from various file formats using file path"""
#     file_extension = file_path.split('.')[-1].lower()
#     text = ''
    
#     if file_extension == 'pdf':
#         # Extract text from PDF using PyMuPDF
#         with open(file_path, 'rb') as file:
#             doc = fitz.open(stream=file.read(), filetype='pdf')
#             for page in doc:
#                 text += page.get_text("text")
#     elif file_extension == 'docx':
#         # Extract text from DOCX
#         doc = Document(file_path)
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#     elif file_extension == 'txt':
#         # Extract text from TXT
#         with open(file_path, 'r', encoding='utf-8') as file:
#             text = file.read()
#     else:
#         raise ValueError('Unsupported file format. Please provide PDF, DOCX, or TXT files.')
    
#     return text

# def extract_text_from_bytes(file_content: bytes, file_extension: str) -> str:
#     """Extract text from various file formats using file bytes"""
#     text = ''
    
#     if file_extension == 'pdf':
#         # Extract text from PDF using PyMuPDF
#         doc = fitz.open(stream=file_content, filetype='pdf')
#         for page in doc:
#             text += page.get_text("text")
#     elif file_extension == 'docx':
#         # Extract text from DOCX
#         doc = Document(io.BytesIO(file_content))
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#     elif file_extension == 'txt':
#         # Extract text from TXT
#         text = file_content.decode('utf-8')
#     else:
#         raise ValueError('Unsupported file format. Please provide PDF, DOCX, or TXT files.')
    
#     return text

# # Excel processing functions
# def load_skill_matrix(file_path: str) -> List[dict]:
#     """Load and process skill matrix Excel file from path"""
#     global sheets_data
#     sheets_data = []
    
#     # Load Excel file
#     xls = pd.ExcelFile(file_path)
#     record_id = 1  # Unique number for each entry
    
#     # Process all sheets from the second one onward
#     for sheet_name in xls.sheet_names[1:]:  # Skip the first sheet
#         df = pd.read_excel(xls, sheet_name=sheet_name)
        
#         # Define custom column names (assuming first 4 columns are fixed)
#         if len(df.columns) >= 4:
#             df.columns = ['First Name', 'Last Name', 'Experience', 'Expertise'] + list(df.columns[4:])
        
#         # Convert each sheet's data to a list of dictionaries with unique numbers
#         records = []
#         for _, row in df.iterrows():
#             record = {"ID": record_id}  # Assign a unique ID
#             record.update(row.to_dict())  # Convert row data to dictionary
#             records.append(record)
#             record_id += 1  # Increment unique ID
        
#         sheet_dict = {
#             "Sheet Name": sheet_name,
#             "Data": records
#         }
        
#         sheets_data.append(sheet_dict)
    
#     return sheets_data

# def load_skill_matrix_from_bytes(file_content: bytes) -> List[dict]:
#     """Load and process skill matrix Excel file from bytes"""
#     global sheets_data
#     sheets_data = []
    
#     # Load Excel file from bytes
#     xls = pd.ExcelFile(io.BytesIO(file_content))
#     record_id = 1  # Unique number for each entry
    
#     # Process all sheets from the second one onward
#     for sheet_name in xls.sheet_names[1:]:  # Skip the first sheet
#         df = pd.read_excel(xls, sheet_name=sheet_name)
        
#         # Define custom column names (assuming first 4 columns are fixed)
#         if len(df.columns) >= 4:
#             df.columns = ['First Name', 'Last Name', 'Experience', 'Expertise'] + list(df.columns[4:])
        
#         # Convert each sheet's data to a list of dictionaries with unique numbers
#         records = []
#         for _, row in df.iterrows():
#             record = {"ID": record_id}  # Assign a unique ID
#             record.update(row.to_dict())  # Convert row data to dictionary
#             records.append(record)
#             record_id += 1  # Increment unique ID
        
#         sheet_dict = {
#             "Sheet Name": sheet_name,
#             "Data": records
#         }
        
#         sheets_data.append(sheet_dict)
    
#     return sheets_data

# def get_json_by_name(first_name: str, last_name: str) -> List[dict]:
#     """Retrieve JSON values based on First Name and Last Name."""
#     results = []
#     for sheet in sheets_data:
#         for record in sheet["Data"]:
#             if (record.get("First Name") == first_name and 
#                 record.get("Last Name") == last_name):
#                 # Include sheet name in record for context
#                 record_with_sheet = record.copy()
#                 record_with_sheet["Sheet Name"] = sheet["Sheet Name"]
#                 results.append(record_with_sheet)
#     return results

# # LLM functions
# def LLM_CALL_1(resume_text: str) -> str:
#     """Extract structured JSON from resume text"""
#     model = "mistral-large-latest"
#     client = Mistral(api_key=API_KEY)
#     parsed_text = resume_text
    
#     # Read system prompt from file
#     system_prompt_path = os.path.join(PROMPTS_DIR, "LLM1.txt")
#     if os.path.exists(system_prompt_path):
#         with open(system_prompt_path, 'r') as file:
#             system_prompt = file.read()
#     else:
#         # Fallback if file doesn't exist
#         system_prompt = """
#         You are an AI assistant that extracts structured information from resumes.
#         Extract relevant details from the resume text and format them into a well-structured JSON object.
#         Include name, designation, objective, education, skills, and project details.
#         """
    
#     chat_response = client.chat.complete(
#         model=model,
#         messages=[
#             {"role": "system", "content": system_prompt},
#             {"role": "user", "content": f"Here is an extracted resume text:\n\n{parsed_text}\n\nFormat it into JSON."}
#         ]
#     )
    
#     json_schema = chat_response.choices[0].message.content
    
#     # Clean the JSON output by removing triple backticks if present
#     cleaned_json = json_schema.strip()
#     if cleaned_json.startswith("```") and cleaned_json.endswith("```"):
#         cleaned_json = cleaned_json.strip("```").strip()
    
#     return cleaned_json

# def LLM_CALL_2(formatted_json_schema: Dict, old_resume_text: str, skill_matrix_json: str) -> str:
#     """Reformat old resume with skill matrix data into new structured JSON"""
#     model = "mistral-large-latest"
#     client = Mistral(api_key=API_KEY)
    
#     system_prompt = """
#     You are an AI assistant that reformats resumes into a structured JSON format.
#     Take the competency matrix and old resume as input, extract relevant details, and map them to the new resume format.
#     Only respond with the new resume format as a JSON object that adheres strictly to the provided JSON Schema. Do not include any extra messages.
#     """
    
#     chat_response = client.chat.complete(
#         model=model,
#         messages=[
#             {"role": "system", "content": system_prompt},
#             {"role": "user", "content": f"""
#             You are given the following details:
            
#             **Skill Matrix (Extracted from External Data Sources):**
#             {skill_matrix_json}
            
#             **Old Resume (Extracted Text):**
#             {old_resume_text}
            
#             Your task is to extract relevant details and format them into the following structured JSON format:
            
#             {json.dumps(formatted_json_schema, indent=2)}
            
#             **Instructions:**
#             - Only use the given details; do not generate fictional information.
#             - Follow the JSON structure exactly as provided.
#             - Ensure the output starts directly with a valid JSON object (no extra text or explanations).
            
#             Generate the updated resume in JSON format:
#             name and designation mandatory key value pairs and create a proper summarized objective a key in for any kind of templete.
#             """}
#         ],
#         response_format = {
#             "type": "json_object",
#         }
#     )
    
#     resume_content = chat_response.choices[0].message.content
#     return resume_content

# def generate_cover_letter_from_resume(resume_summary: Dict) -> str:
#     """Generates a professional cover letter based on the candidate's resume summary"""
#     model = "mistral-large-latest"
#     client = Mistral(api_key=API_KEY)
    
#     system_prompt = """
#     You are an AI assistant that creates professional cover letters based only on a candidate's resume summary.
#     Your task is to analyze the resume summary, infer the candidate's expertise, and generate a well-structured cover letter.
#     Ensure the tone is formal, engaging, and professional.
#     """
    
#     user_prompt = f"""
#     Candidate Resume Summary:
#     {json.dumps(resume_summary, indent=2)}
    
#     Based on this summary, write a professional cover letter.
#     Follow this structure:
#     - Address the hiring manager (use "Dear Hiring Manager" if no specific name is provided).
#     - Introduce the candidate and express general interest in roles that match their expertise.
#     - Highlight key experiences, achievements, and skills relevant to their domain.
#     - Conclude with enthusiasm and a call to action.
    
#     Ensure the letter is formal and engaging.
#     """
    
#     chat_response = client.chat.complete(
#         model=model,
#         messages=[
#             {"role": "system", "content": system_prompt},
#             {"role": "user", "content": user_prompt}
#         ],
#         response_format={"type": "text"}
#     )
    
#     cover_letter = chat_response.choices[0].message.content
#     return cover_letter

# def generate_resume_1(data_json: str) -> str:
#     """Generates the resume PDF with precise formatting."""
#     data = json.loads(data_json) if isinstance(data_json, str) else data_json
#     pdf = PDF()
#     pdf.set_margins(15, 15, 15)  # Set consistent margins
#     pdf.add_page()

#     # Header Information
#     pdf.set_font("Arial", "B", 16)
#     pdf.cell(0, 8, data["name"], ln=True, align="L")

#     pdf.set_font("Arial", "I", 12)
#     pdf.cell(0, 8, data["designation"], ln=True, align="L")  # Added Designation Field

#     # **Horizontal Line**
#     pdf.set_line_width(0.5)  # Line thickness
#     pdf.line(15, pdf.get_y(), 195, pdf.get_y())  # Draws a line across the page
#     pdf.ln(8)

#     # Sections
#     pdf.chapter_title("PROFESSIONAL SUMMARY")
#     pdf.set_font("Arial", "", 9)
#     pdf.multi_cell(0, 7, data["objective"])
#     pdf.ln(5)

#     # Education Section
#     pdf.chapter_title("EDUCATION")
#     pdf.multi_cell(0, 7, "\n".join(data["education"]))  # List all education details
#     pdf.ln(5)

#     # Skills Section
#     pdf.chapter_title("SKILLS")
#     pdf.set_font("Arial", "", 9)  # Reduce text size and set normal font
#     pdf.multi_cell(0, 6, ", ".join(data["skills"]))  # Reduce line height
#     pdf.ln(8)

#     # Projects Section
#     pdf.chapter_title("PROJECT DETAILS")
#     for project in data["project_details"].values():  # Iterate over dictionary values
#         pdf.add_project_table(project)

#     # Save PDF
#     filename = os.path.join(OUTPUT_DIR, f"{data['name'].replace(' ', '_')}_Resume.pdf")
#     pdf.output(filename)
#     return filename

# def generate_cover_letter_pdf(cover_letter_text: str, candidate_name: str) -> str:
#     """Generates a PDF file from the cover letter text."""
#     pdf = PDF()
#     pdf.set_margins(15, 15, 15)
#     pdf.add_page()
    
#     # Header with candidate name
#     pdf.set_font("Arial", "B", 16)
#     pdf.cell(0, 8, f"{candidate_name} - Cover Letter", ln=True, align="L")
    
#     # Horizontal Line
#     pdf.set_line_width(0.5)
#     pdf.line(15, pdf.get_y(), 195, pdf.get_y())
#     pdf.ln(8)
    
#     # Cover letter content
#     pdf.set_font("Arial", "", 11)
    
#     # Split text into paragraphs and add to PDF
#     paragraphs = cover_letter_text.split("\n\n")
#     for paragraph in paragraphs:
#         if paragraph.strip():
#             pdf.multi_cell(0, 7, paragraph.strip())
#             pdf.ln(4)
    
#     # Save the cover letter
#     filename = os.path.join(OUTPUT_DIR, f"{candidate_name.replace(' ', '_')}_Cover_Letter.pdf")
#     pdf.output(filename)
#     return filename

# # FastAPI endpoints
# @app.post("/upload/resume", response_model=dict)
# async def upload_resume(file: UploadFile = File(...)):
#     """
#     Upload a resume file (PDF, DOCX, or TXT) and extract text from it.
#     """
#     try:
#         file_extension = file.filename.split('.')[-1].lower()
#         if file_extension not in ['pdf', 'docx', 'txt']:
#             raise HTTPException(status_code=400, detail="Unsupported file format. Please provide PDF, DOCX, or TXT files.")
        
#         file_content = await file.read()
#         extracted_text = extract_text_from_bytes(file_content, file_extension)
        
#         # Save the file temporarily
#         temp_file_path = os.path.join(OUTPUT_DIR, file.filename)
#         with open(temp_file_path, "wb") as temp_file:
#             with open(temp_file_path, "wb") as f:
#                 f.write(file_content)
        
#         return {
#             "filename": file.filename,
#             "text_content": extracted_text,
#             "file_path": temp_file_path
#         }
    
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

# @app.post("/upload/skill-matrix", response_model=dict)
# async def upload_skill_matrix(file: UploadFile = File(...)):
#     """
#     Upload an Excel file containing skill matrix data and process it.
#     """
#     try:
#         file_extension = file.filename.split('.')[-1].lower()
#         if file_extension not in ['xlsx', 'xls']:
#             raise HTTPException(status_code=400, detail="Please upload an Excel file (.xlsx, .xls)")
        
#         file_content = await file.read()
#         processed_data = load_skill_matrix_from_bytes(file_content)
        
#         # Save the file temporarily
#         temp_file_path = os.path.join(OUTPUT_DIR, file.filename)
#         with open(temp_file_path, "wb") as f:
#             f.write(file_content)
        
#         return {
#             "filename": file.filename,
#             "sheets_processed": len(processed_data),
#             "total_records": sum(len(sheet["Data"]) for sheet in processed_data),
#             "file_path": temp_file_path
#         }
    
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Error processing Excel file: {str(e)}")

# @app.post("/search/by-name", response_model=List[dict])
# async def search_by_name(query: NameQuery):
#     """
#     Search for skill matrix records by first name and last name.
#     """
#     if not sheets_data:
#         raise HTTPException(status_code=400, 
#                            detail="No skill matrix data loaded. Please upload a skill matrix Excel file first.")
    
#     results = get_json_by_name(query.first_name, query.last_name)
    
#     if not results:
#         return []
    
#     return results

# @app.post("/generate/resume", response_model=dict)
# async def generate_resume(request: ResumeRequest, background_tasks: BackgroundTasks):
#     """
#     Generate a new resume based on old resume, template, and skill matrix data.
#     """
#     try:
#         # Extract text from old resume
#         old_resume_text = extract_text_from_file(request.old_resume_path)
        
#         # Extract text from the template
#         template_text = extract_text_from_file(request.template_path)
        
#         # Get schema from template
#         json_schema = LLM_CALL_1(template_text)
        
#         # Parse JSON string into a Python dictionary
#         try:
#             structured_data = json.loads(json_schema)
#         except json.JSONDecodeError:
#             # Try cleaning the JSON if initial parsing fails
#             cleaned_json = json_schema.strip()
#             if cleaned_json.startswith("```json") and cleaned_json.endswith("```"):
#                 cleaned_json = cleaned_json[7:-3].strip()  # Remove ```json and ``` markers
#             elif cleaned_json.startswith("```") and cleaned_json.endswith("```"):
#                 cleaned_json = cleaned_json[3:-3].strip()  # Remove ``` markers
#             structured_data = json.loads(cleaned_json)
        
#         # Get skill matrix data if provided
#         skill_matrix_json = "{}"
#         if request.skill_matrix_path:
#             load_skill_matrix(request.skill_matrix_path)
#             if request.first_name and request.last_name:
#                 skill_data = get_json_by_name(request.first_name, request.last_name)
#                 if skill_data:
#                     skill_matrix_json = json.dumps(skill_data)
        
#         # Generate new resume content
#         new_resume_content = LLM_CALL_2(structured_data, old_resume_text, skill_matrix_json)
#         new_resume_content_json = json.loads(new_resume_content)
        
#         # Generate PDF resume
#         resume_path = generate_resume_1(new_resume_content_json)
        
#         # Generate cover letter in background
#         background_tasks.add_task(
#             generate_and_save_cover_letter, 
#             new_resume_content_json, 
#             new_resume_content_json["name"]
#         )
        
#         return {
#             "message": "Resume generation completed successfully",
#             "resume_path": resume_path,
#             "cover_letter_status": "Processing in background",
#             "resume_json": new_resume_content_json
#         }
    
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"Error generating resume: {str(e)}")

# def generate_and_save_cover_letter(resume_json: Dict, candidate_name: str):
#     """Background task to generate and save a cover letter"""
#     try:
#         # Generate cover letter text
#         cover_letter_text = generate_cover_letter_from_resume(resume_json)
        
#         # Generate cover letter PDF
#         cover_letter_path = generate_cover_letter_pdf(cover_letter_text, candidate_name)
        
#         # Save text version
#         text_path = os.path.join(OUTPUT_DIR, f"{candidate_name.replace(' ', '_')}_Cover_Letter.txt")
#         with open(text_path, "w") as f:
#             f.write(cover_letter_text)
            
#         print(f"Cover letter generated successfully: {cover_letter_path}")
#     except Exception as e:
#         print(f"Error generating cover letter: {str(e)}")

# @app.get("/download/{file_path:path}")
# async def download_file(file_path: str):
#     """
#     Download a generated file.
#     """
#     full_path = os.path.join(OUTPUT_DIR, file_path)
#     if not os.path.exists(full_path):
#         raise HTTPException(status_code=404, detail="File not found")
    
#     return FileResponse(path=full_path, filename=os.path.basename(full_path))

# @app.get("/status")
# async def check_status():
#     """
#     Check API status.
#     """
#     return {"status": "active", "message": "Resume Automation API is running"}

# if __name__ == "__main__":
#     import uvicorn
    
#     # Ensure directories exist
#     os.makedirs(OUTPUT_DIR, exist_ok=True)
#     os.makedirs(PROMPTS_DIR, exist_ok=True)
    
#     # Run the FastAPI application
#     uvicorn.run(app, host="0.0.0.0", port=8000)











# import uvicorn
# from fastapi import FastAPI
# from fastapi.staticfiles import StaticFiles
# from setup import STATIC_DIR, TEMPLATES_DIR
# from routes import router

# app = FastAPI()

# # Mount API routes
# app.include_router(router)

# # Serve static files
# app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

# @app.get("/")
# async def root():
#     """Return a simple welcome message"""
#     return {"message": "Welcome to Resume Automation API"}

# if __name__ == "__main__":
#     uvicorn.run(app, host="0.0.0.0", port=8000)






from fastapi import FastAPI, File, UploadFile, Form, HTTPException, BackgroundTasks
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional, Dict, Any
import pandas as pd
import json
import fitz  # PyMuPDF
from docx import Document
import io
import os
from pydantic import BaseModel
from mistralai import Mistral
from fpdf import FPDF
from fpdf import FPDF
# import json
import re

app = FastAPI(title="Resume Automation API")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# API key for Mistral
API_KEY = "ZJjgVeyKcDyuN7PGriE6mHR6kKsN8gdq"

# Global variable to store loaded skill matrix data
sheets_data = []

# Paths configuration
BASE_PATH = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(BASE_PATH, "Logo.png")
PROMPTS_DIR = os.path.join(BASE_PATH, "Prompts")
OUTPUT_DIR = os.path.join(BASE_PATH, "Output")
TEMPLATES_DIR = os.path.join(BASE_PATH, "templates")

# Create required directories if they don't exist
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(PROMPTS_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)

# Create HTML template directory and file
with open(os.path.join(TEMPLATES_DIR, "index.html"), "w") as f:
    f.write("""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume Automation Tool</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }
        h1 {
            color: #2c3e50;
            text-align: center;
            margin-bottom: 30px;
        }
        .container {
            display: flex;
            flex-direction: column;
            gap: 20px;
        }
        .card {
            border: 1px solid #e0e0e0;
            border-radius: 5px;
            padding: 20px;
            box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);
        }
        .card h2 {
            margin-top: 0;
            color: #3498db;
            border-bottom: 1px solid #eee;
            padding-bottom: 10px;
        }
        form {
            display: flex;
            flex-direction: column;
            gap: 15px;
        }
        label {
            font-weight: bold;
        }
        input, button {
            padding: 8px 12px;
            border-radius: 4px;
            border: 1px solid #ddd;
        }
        input[type="file"] {
            border: none;
        }
        button {
            background-color: #3498db;
            color: white;
            border: none;
            cursor: pointer;
            font-weight: bold;
            transition: background-color 0.3s;
        }
        button:hover {
            background-color: #2980b9;
        }
        #result {
            margin-top: 20px;
            padding: 15px;
            border-radius: 5px;
            background-color: #f8f9fa;
            border-left: 5px solid #28a745;
            display: none;
        }
        #search-results {
            margin-top: 15px;
        }
        .hidden {
            display: none;
        }
        .file-info {
            margin-top: 10px;
            background-color: #e8f4fd;
            padding: 10px;
            border-radius: 4px;
            display: none;
        }
        .steps {
            display: flex;
            justify-content: space-between;
            margin-bottom: 20px;
        }
        .step {
            background-color: #f1f1f1;
            padding: 10px 15px;
            border-radius: 20px;
            font-weight: bold;
        }
        .step.active {
            background-color: #3498db;
            color: white;
        }
    </style>
</head>
<body>
    <h1>Resume Automation Tool</h1>
    
    <div class="container">
        <div class="card">
            <h2>1. Upload Skill Matrix</h2>
            <form id="skill-matrix-form">
                <label for="skill-matrix">Upload Excel Skill Matrix File:</label>
                <input type="file" id="skill-matrix" accept=".xlsx,.xls" required>
                <button type="submit">Upload Skill Matrix</button>
            </form>
            <div id="skill-matrix-info" class="file-info"></div>
        </div>

        <div class="card">
            <h2>2. Upload Resume</h2>
            <form id="resume-form">
                <label for="resume">Upload Resume File (PDF, DOCX, TXT):</label>
                <input type="file" id="resume" accept=".pdf,.docx,.txt" required>
                <button type="submit">Upload Resume</button>
            </form>
            <div id="resume-info" class="file-info"></div>
        </div>

        <div class="card">
            <h2>3. Upload Resume Template</h2>
            <form id="template-form">
                <label for="template">Upload Template File (PDF, DOCX, TXT):</label>
                <input type="file" id="template" accept=".pdf,.docx,.txt" required>
                <button type="submit">Upload Template</button>
            </form>
            <div id="template-info" class="file-info"></div>
        </div>

        <div class="card">
            <h2>4. Search Employee in Skill Matrix</h2>
            <form id="search-form">
                <label for="first-name">First Name:</label>
                <input type="text" id="first-name" required>
                <label for="last-name">Last Name:</label>
                <input type="text" id="last-name" required>
                <button type="submit">Search</button>
            </form>
            <div id="search-results"></div>
        </div>

        <div class="card">
            <h2>5. Generate Resume</h2>
            <form id="generate-form">
                <p>Using the files uploaded above and search results, generate a new resume.</p>
                <button type="submit" id="generate-btn">Generate Resume</button>
            </form>
            <div id="result"></div>
        </div>
    </div>

    <script>
        // Store file paths
        const filePaths = {
            skillMatrix: '',
            resume: '',
            template: ''
        };
        
        let firstName = '';
        let lastName = '';

        // Skill Matrix Upload
        document.getElementById('skill-matrix-form').addEventListener('submit', async function(e) {
            e.preventDefault();
            const fileInput = document.getElementById('skill-matrix');
            const file = fileInput.files[0];
            if (!file) return;

            const formData = new FormData();
            formData.append('file', file);

            try {
                const response = await fetch('/upload/skill-matrix', {
                    method: 'POST',
                    body: formData
                });

                const data = await response.json();
                if (response.ok) {
                    filePaths.skillMatrix = data.file_path;
                    document.getElementById('skill-matrix-info').textContent = 
                        `File uploaded: ${data.filename}. Processed ${data.sheets_processed} sheets with ${data.total_records} records.`;
                    document.getElementById('skill-matrix-info').style.display = 'block';
                } else {
                    alert('Error: ' + data.detail);
                }
            } catch (error) {
                console.error('Error:', error);
                alert('Failed to upload file');
            }
        });

        // Resume Upload
        document.getElementById('resume-form').addEventListener('submit', async function(e) {
            e.preventDefault();
            const fileInput = document.getElementById('resume');
            const file = fileInput.files[0];
            if (!file) return;

            const formData = new FormData();
            formData.append('file', file);

            try {
                const response = await fetch('/upload/resume', {
                    method: 'POST',
                    body: formData
                });

                const data = await response.json();
                if (response.ok) {
                    filePaths.resume = data.file_path;
                    document.getElementById('resume-info').textContent = 
                        `File uploaded: ${data.filename}. Text extracted successfully.`;
                    document.getElementById('resume-info').style.display = 'block';
                } else {
                    alert('Error: ' + data.detail);
                }
            } catch (error) {
                console.error('Error:', error);
                alert('Failed to upload file');
            }
        });

        // Template Upload
        document.getElementById('template-form').addEventListener('submit', async function(e) {
            e.preventDefault();
            const fileInput = document.getElementById('template');
            const file = fileInput.files[0];
            if (!file) return;

            const formData = new FormData();
            formData.append('file', file);

            try {
                const response = await fetch('/upload/resume', {
                    method: 'POST',
                    body: formData
                });

                const data = await response.json();
                if (response.ok) {
                    filePaths.template = data.file_path;
                    document.getElementById('template-info').textContent = 
                        `File uploaded: ${data.filename}. Text extracted successfully.`;
                    document.getElementById('template-info').style.display = 'block';
                } else {
                    alert('Error: ' + data.detail);
                }
            } catch (error) {
                console.error('Error:', error);
                alert('Failed to upload file');
            }
        });

        // Search Form
      // Search Form
document.getElementById('search-form').addEventListener('submit', async function(e) {
    e.preventDefault();
    firstName = document.getElementById('first-name').value;
    lastName = document.getElementById('last-name').value;

    if (!firstName || !lastName) {
        alert('Please enter both first and last name');
        return;
    }

    try {
        const response = await fetch('/search/by-name', {
            method: 'POST',
            headers: {
                'Content-Type': 'application/json'
            },
            body: JSON.stringify({
                first_name: firstName,
                last_name: lastName
            })
        });

        const data = await response.json();
        const resultsDiv = document.getElementById('search-results');
        
        if (data.length > 0) {
            let html = '<h3>Search Results:</h3><ul>';
            data.forEach(item => {
                // Handle different possible column names
                const id = item.ID || item.Id || "N/A";
                const firstName = item.First_Name || item["First Name"] || "N/A";
                const lastName = item.Last_Name || item["Last Name"] || "N/A";
                const experience = item.Experience || "N/A";
                const sheetName = item["Sheet Name"] || "Unknown";
                
                html += `<li>ID: ${id}, Name: ${firstName} ${lastName}, 
                        Experience: ${experience}, 
                        Sheet: ${sheetName}</li>`;
            });
            html += '</ul>';
            resultsDiv.innerHTML = html;
        } else {
            resultsDiv.innerHTML = '<p>No results found. Please check the name and try again.</p>';
        }
    } catch (error) {
        console.error('Error:', error);
        alert('Failed to search');
    }
});

        // Generate Resume
        // Generate Resume
document.getElementById('generate-form').addEventListener('submit', async function(e) {
    e.preventDefault();
    
    // Show loading state
    const generateBtn = document.getElementById('generate-btn');
    const originalBtnText = generateBtn.textContent;
    generateBtn.textContent = 'Generating...';
    generateBtn.disabled = true;
    
    // Check if all required files are uploaded
    if (!filePaths.resume || !filePaths.template) {
        alert('Please upload resume and template files first');
        generateBtn.textContent = originalBtnText;
        generateBtn.disabled = false;
        return;
    }

    const resultDiv = document.getElementById('result');
    resultDiv.innerHTML = '<p>Processing your request... This may take up to 1 minute.</p>';
    resultDiv.style.display = 'block';

    try {
        const response = await fetch('/generate/resume', {
            method: 'POST',
            headers: {
                'Content-Type': 'application/json'
            },
            body: JSON.stringify({
                old_resume_path: filePaths.resume,
                template_path: filePaths.template,
                skill_matrix_path: filePaths.skillMatrix || null,
                first_name: firstName || null,
                last_name: lastName || null
            })
        });

        const data = await response.json();
        
        if (response.ok) {
            resultDiv.innerHTML = `
                <h3>Resume Generation Complete!</h3>
                <p>${data.message}</p>
                <p>Download files:</p>
                <ul>
                    <li><a href="/download/${data.resume_path.split('/').pop()}" target="_blank">Download Resume</a></li>
                    <li>Cover Letter: ${data.cover_letter_status}</li>
                </ul>
                <p>Check the Output directory for your files.</p>
            `;
        } else {
            resultDiv.innerHTML = `
                <h3>Error Generating Resume</h3>
                <p>Error: ${data.detail || 'Unknown error occurred'}</p>
                <p>Please try again or check the following:</p>
                <ul>
                    <li>Make sure your resume and template files are valid</li>
                    <li>Verify that the skill matrix contains the searched name</li>
                    <li>Try with different input files if the problem persists</li>
                </ul>
            `;
        }
    } catch (error) {
        console.error('Error:', error);
        resultDiv.innerHTML = `
            <h3>Request Failed</h3>
            <p>There was a problem communicating with the server. Please try again later.</p>
            <p>Error details: ${error.message || 'Unknown error'}</p>
        `;
    } finally {
        // Restore button state
        generateBtn.textContent = originalBtnText;
        generateBtn.disabled = false;
    }
});
    </script>
</body>
</html>
    """)

# Setup templates
templates = Jinja2Templates(directory=TEMPLATES_DIR)

# Pydantic models
class NameQuery(BaseModel):
    first_name: str
    last_name: str

class ResumeRequest(BaseModel):
    old_resume_path: str
    template_path: str
    skill_matrix_path: Optional[str] = None
    first_name: Optional[str] = None
    last_name: Optional[str] = None


# File extraction functions
def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file formats using file path"""
    file_extension = file_path.split('.')[-1].lower()
    text = ''
    
    if file_extension == 'pdf':
        # Extract text from PDF using PyMuPDF
        with open(file_path, 'rb') as file:
            doc = fitz.open(stream=file.read(), filetype='pdf')
            for page in doc:
                text += page.get_text("text")
    elif file_extension == 'docx':
        # Extract text from DOCX
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    elif file_extension == 'txt':
        # Extract text from TXT
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    else:
        raise ValueError('Unsupported file format. Please provide PDF, DOCX, or TXT files.')
    
    return text

def extract_text_from_bytes(file_content: bytes, file_extension: str) -> str:
    """Extract text from various file formats using file bytes"""
    text = ''
    
    if file_extension == 'pdf':
        # Extract text from PDF using PyMuPDF
        doc = fitz.open(stream=file_content, filetype='pdf')
        for page in doc:
            text += page.get_text("text")
    elif file_extension == 'docx':
        # Extract text from DOCX
        doc = Document(io.BytesIO(file_content))
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    elif file_extension == 'txt':
        # Extract text from TXT
        text = file_content.decode('utf-8')
    else:
        raise ValueError('Unsupported file format. Please provide PDF, DOCX, or TXT files.')
    
    return text

# Excel processing functions
def load_skill_matrix(file_path: str) -> List[dict]:
    """Load and process skill matrix Excel file from path"""
    global sheets_data
    sheets_data = []
    
    # Load Excel file
    xls = pd.ExcelFile(file_path)
    record_id = 1  # Unique number for each entry
    
    # Process all sheets from the second one onward
    for sheet_name in xls.sheet_names[1:]:  # Skip the first sheet
        df = pd.read_excel(xls, sheet_name=sheet_name)
        
        # Standardize column names to ensure consistency
        if len(df.columns) >= 4:
            # Create standardized column names
            std_columns = ['First_Name', 'Last_Name', 'Experience', 'Expertise'] + list(df.columns[4:])
            
            # Map original columns to standardized names
            column_mapping = {original: standard for original, standard in zip(df.columns, std_columns)}
            df = df.rename(columns=column_mapping)
        
        # Convert each sheet's data to a list of dictionaries with unique numbers
        records = []
        for _, row in df.iterrows():
            record = {"ID": record_id}  # Assign a unique ID
            record.update(row.to_dict())  # Convert row data to dictionary
            records.append(record)
            record_id += 1  # Increment unique ID
        
        sheet_dict = {
            "Sheet Name": sheet_name,
            "Data": records
        }
        
        sheets_data.append(sheet_dict)
    
    return sheets_data

def load_skill_matrix_from_bytes(file_content: bytes) -> List[dict]:
    """Load and process skill matrix Excel file from bytes"""
    global sheets_data
    sheets_data = []
    
    # Load Excel file from bytes
    xls = pd.ExcelFile(io.BytesIO(file_content))
    record_id = 1  # Unique number for each entry
    
    # Process all sheets from the second one onward
    for sheet_name in xls.sheet_names[1:]:  # Skip the first sheet
        df = pd.read_excel(xls, sheet_name=sheet_name)
        
        # Standardize column names to ensure consistency
        if len(df.columns) >= 4:
            # Create standardized column names
            std_columns = ['First_Name', 'Last_Name', 'Experience', 'Expertise'] + list(df.columns[4:])
            
            # Map original columns to standardized names
            column_mapping = {original: standard for original, standard in zip(df.columns, std_columns)}
            df = df.rename(columns=column_mapping)
        
        # Convert each sheet's data to a list of dictionaries with unique numbers
        records = []
        for _, row in df.iterrows():
            record = {"ID": record_id}  # Assign a unique ID
            record.update(row.to_dict())  # Convert row data to dictionary
            records.append(record)
            record_id += 1  # Increment unique ID
        
        sheet_dict = {
            "Sheet Name": sheet_name,
            "Data": records
        }
        
        sheets_data.append(sheet_dict)
    
    return sheets_data

def get_json_by_name(first_name: str, last_name: str) -> List[dict]:
    """Retrieve JSON values based on First Name and Last Name."""
    results = []
    for sheet in sheets_data:
        for record in sheet["Data"]:
            # Check for both standardized and possible original column names
            fn_match = (record.get("First_Name") == first_name or record.get("First Name") == first_name)
            ln_match = (record.get("Last_Name") == last_name or record.get("Last Name") == last_name)
            
            if fn_match and ln_match:
                # Include sheet name in record for context
                record_with_sheet = record.copy()
                record_with_sheet["Sheet Name"] = sheet["Sheet Name"]
                results.append(record_with_sheet)
    return results


# LLM functions
def LLM_CALL_1(resume_text: str) -> str:
    """Extract structured JSON from resume text"""
    try:
        model = "mistral-large-latest"
        client = Mistral(api_key=API_KEY)
        parsed_text = resume_text[:5000]  # Limit text length to avoid token limits
        
        # Read system prompt from file
        system_prompt_path = os.path.join(PROMPTS_DIR, "LLM1.txt")
        if os.path.exists(system_prompt_path):
            with open(system_prompt_path, 'r') as file:
                system_prompt = file.read()
        else:
            # Fallback if file doesn't exist
            system_prompt = """
            You are an AI assistant that extracts structured information from resumes.
            Extract relevant details from the resume text and format them into a well-structured JSON object.
            Include name, designation, objective, education, skills, and project details.
            Your response should be a valid JSON object without any additional text or markdown formatting.
            """
        
        try:
            chat_response = client.chat.complete(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Here is an extracted resume text:\n\n{parsed_text}\n\nFormat it into JSON."}
                ],
                response_format = {
                    "type": "json_object",
                }
            )
            
            json_schema = chat_response.choices[0].message.content
            
            # Clean the JSON output by removing triple backticks if present
            cleaned_json = json_schema.strip()
            if cleaned_json.startswith("```json") and cleaned_json.endswith("```"):
                cleaned_json = cleaned_json[7:-3].strip()
            elif cleaned_json.startswith("```") and cleaned_json.endswith("```"):
                cleaned_json = cleaned_json[3:-3].strip()
            
            # Validate the JSON
            try:
                json.loads(cleaned_json)
                return cleaned_json
            except json.JSONDecodeError:
                # If invalid JSON, try to fix common issues
                import re
                cleaned_json = cleaned_json.replace('\n', ' ').replace('\r', '')
                cleaned_json = re.sub(r'(?<!")(\w+)(?=":)', r'"\1"', cleaned_json)
                return cleaned_json
                
        except Exception as e:
            print(f"Error in Mistral API call: {str(e)}")
            # Return a default JSON structure
            default_json = {
                "name": "",
                "designation": "",
                "objective": "",
                "education": [],
                "skills": [],
                "project_details": {}
            }
            return json.dumps(default_json)
    
    except Exception as e:
        print(f"Unexpected error in LLM_CALL_1: {str(e)}")
        default_json = {
            "name": "",
            "designation": "",
            "objective": "",
            "education": [],
            "skills": [],
            "project_details": {}
        }
        return json.dumps(default_json)

def LLM_CALL_2(formatted_json_schema: Dict, old_resume_text: str, skill_matrix_json: str) -> str:
    """Reformat old resume with skill matrix data into new structured JSON"""
    try:
        model = "mistral-large-latest"
        client = Mistral(api_key=API_KEY)
        
        # Limit text lengths to avoid token limits
        old_resume_text = old_resume_text[:4000]
        if len(skill_matrix_json) > 2000:
            skill_matrix_json = skill_matrix_json[:2000] + "... [truncated]"
        
        system_prompt = """
        You are an AI assistant that reformats resumes into a structured JSON format.
        Take the competency matrix and old resume as input, extract relevant details, and map them to the new resume format.
        Only respond with the new resume format as a JSON object that adheres strictly to the provided JSON Schema.
        Your response must be a valid, properly formatted JSON object without any additional text, explanations, or markdown formatting.
        """
        
        try:
            chat_response = client.chat.complete(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"""
                    You are given the following details:
                    
                    **Skill Matrix (Extracted from External Data Sources):**
                    {skill_matrix_json}
                    
                    **Old Resume (Extracted Text):**
                    {old_resume_text}
                    
                    Your task is to extract relevant details and format them into the following structured JSON format:
                    
                    {json.dumps(formatted_json_schema, indent=2)}
                    
                    **Instructions:**
                    - Only use the given details; do not generate fictional information.
                    - Follow the JSON structure exactly as provided.
                    - Ensure the output is a valid JSON object (no extra text or explanations).
                    - "name" and "designation" are mandatory fields.
                    - Create a proper summarized objective.
                    
                    Generate the updated resume in JSON format:
                    """}
                ],
                response_format = {
                    "type": "json_object",
                }
            )
            
            resume_content = chat_response.choices[0].message.content
            
            # Attempt to validate and clean JSON if needed
            try:
                json.loads(resume_content)
                return resume_content
            except json.JSONDecodeError:
                # Clean up if it's not valid JSON
                cleaned_content = resume_content.replace('```json', '').replace('```', '').strip()
                return cleaned_content
                
        except Exception as e:
            print(f"Error in Mistral API call: {str(e)}")
            # Extract name from resume text as fallback
            name = "Candidate"
            for line in old_resume_text.split('\n'):
                if line.strip() and len(line.strip()) < 50:
                    name = line.strip()
                    break
                    
            # Return a basic JSON with resume text information
            default_json = {
                "name": name,
                "designation": "Professional",
                "objective": "Experienced professional seeking new opportunities.",
                "education": ["Education details not extracted"],
                "skills": ["Skills not extracted"],
                "project_details": {
                    "project1": {
                        "name": "Project",
                        "role": "Team Member",
                        "description": "Project description not extracted",
                        "technology": "Technologies not extracted",
                        "role_played": "Role details not extracted"
                    }
                }
            }
            return json.dumps(default_json)
    
    except Exception as e:
        print(f"Unexpected error in LLM_CALL_2: {str(e)}")
        # Return a basic JSON structure
        default_json = {
            "name": "Candidate",
            "designation": "Professional",
            "objective": "Experienced professional seeking new opportunities.",
            "education": ["Education details not extracted"],
            "skills": ["Skills not extracted"],
            "project_details": {
                "project1": {
                    "name": "Project",
                    "role": "Team Member",
                    "description": "Project description not extracted",
                    "technology": "Technologies not extracted",
                    "role_played": "Role details not extracted"
                }
            }
        }
        return json.dumps(default_json)
    

def generate_cover_letter_from_resume(resume_summary: Dict) -> str:
    """Generates a professional cover letter based on the candidate's resume summary"""
    model = "mistral-large-latest"
    client = Mistral(api_key=API_KEY)
    
    system_prompt = """
    You are an AI assistant that creates professional cover letters based only on a candidate's resume summary.
    Your task is to analyze the resume summary, infer the candidate's expertise, and generate a well-structured cover letter.
    Ensure the tone is formal, engaging, and professional.
    """
    
    user_prompt = f"""
    Candidate Resume Summary:
    {json.dumps(resume_summary, indent=2)}
    
    Based on this summary, write a professional cover letter.
    Follow this structure:
    - Address the hiring manager (use "Dear Hiring Manager" if no specific name is provided).
    - Introduce the candidate and express general interest in roles that match their expertise.
    - Highlight key experiences, achievements, and skills relevant to their domain.
    - Conclude with enthusiasm and a call to action.
    
    Ensure the letter is formal and engaging.
    """
    
    chat_response = client.chat.complete(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        response_format={"type": "text"}
    )
    
    cover_letter = chat_response.choices[0].message.content
    return cover_letter






class PDF(FPDF):
    def header(self):
        if hasattr(self, 'logo_path') and self.logo_path:
            self.image(self.logo_path, 160, 8, 35)
        self.set_font("Arial", "B", 12)
        self.ln(10)

    def chapter_title(self, title):
        self.set_text_color(0, 0, 180)
        self.set_font("Arial", "B", 11)
        self.cell(0, 8, self.clean_text(title), ln=True, align="L")
        self.ln(4)
        self.set_text_color(0, 0, 0)

    def add_project_table(self, project):
        """Creates a vertical-style project details table."""
        col_widths = [50, 120]
        row_height = 8

        self.set_font("Arial", "B", 9)
        self.cell(0, 6, f"<{self.clean_text(project.get('project_name', 'N/A'))}>", ln=True)
        self.ln(2)
        self.set_font("Arial", "", 9)

        def add_row(label, text):
            if text:
                text = self.clean_text(text)
                text_height = (self.get_string_width(text) // col_widths[1] + 1) * row_height
                self.cell(col_widths[0], text_height, label, border=1, align="C")
                self.multi_cell(col_widths[1], row_height, text, border=1)

        fields = [
            ("Project Name", project.get("project_name", "")),
            ("Role", project.get("role", "")),
            ("Description", project.get("description", "")),
            ("Technology", project.get("technology", "")),
            ("Role Played", project.get("role_played", "")),
        ]

        for label, text in fields:
            add_row(label, text)

        self.ln()


    def cell(self, w, h=0, txt='', border=0, ln=0, align='', fill=False, link=''):
        txt = self.clean_text(txt)
        super().cell(w, h, txt, border, ln, align, fill, link)

    def multi_cell(self, w, h, txt, border=0, align='J', fill=False):
        txt = self.clean_text(txt)
        super().multi_cell(w, h, txt, border, align, fill)

    def clean_text(self, text):
        if not isinstance(text, str):
            text = str(text)
        text = re.sub(r"[“”]", '"', text)
        text = re.sub(r"[‘’]", "'", text)
        return text.encode("latin-1", "replace").decode("latin-1")
def generate_resume_1(data_json: str) -> str:
    """Generates the resume PDF with precise formatting."""
    # Handle both string and dictionary input types
    if isinstance(data_json, str):
        data = json.loads(data_json)
    else:
        data = data_json
    
    pdf = PDF()
    pdf.set_margins(15, 15, 15)  # Set consistent margins
    pdf.add_page()

    # Header Information
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 8, data["name"], ln=True, align="L")

    pdf.set_font("Arial", "I", 12)
    pdf.cell(0, 8, data["designation"], ln=True, align="L")

    # Horizontal Line
    pdf.set_line_width(0.5)  # Line thickness
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())  # Draws a line across the page
    pdf.ln(8)

    # Sections
    pdf.chapter_title("PROFESSIONAL SUMMARY")
    pdf.set_font("Arial", "", 9)
    # Use 'objective' instead of 'summarized_objective'
    pdf.multi_cell(0, 7, data["objective"])
    pdf.ln(5)

    # Education Section
    pdf.chapter_title("EDUCATION")
    pdf.multi_cell(0, 7, "\n".join(data["education"]))  # List all education details
    pdf.ln(5)

    # Skills Section
    pdf.chapter_title("SKILLS")
    pdf.set_font("Arial", "", 9)  # Reduce text size and set normal font
    pdf.multi_cell(0, 6, ", ".join(data["skills"]))  # Reduce line height
    pdf.ln(8)

    # Projects Section
    pdf.chapter_title("PROJECT DETAILS")
    # Handle project_details as a dictionary or a list
    if isinstance(data["project_details"], dict):
        for project in data["project_details"].values():  # Iterate over dictionary values
            pdf.add_project_table(project)
    elif isinstance(data["project_details"], list):
        for project in data["project_details"]:  # Iterate over list
            pdf.add_project_table(project)

    # Save PDF
    filename = os.path.join(OUTPUT_DIR, f"{data['name'].replace(' ', '_')}_Resume.pdf")
    pdf.output(filename)
    return filename



def generate_cover_letter_pdf(cover_letter_text: str, candidate_name: str) -> str:
    """Generates a PDF file from the cover letter text."""
    pdf = PDF()
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    
    # Header with candidate name
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 8, f"{candidate_name} - Cover Letter", ln=True, align="L")
    
    # Horizontal Line
    pdf.set_line_width(0.5)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(8)
    
    # Cover letter content
    pdf.set_font("Arial", "", 11)
    
    # Split text into paragraphs and add to PDF
    paragraphs = cover_letter_text.split("\n\n")
    for paragraph in paragraphs:
        if paragraph.strip():
            pdf.multi_cell(0, 7, paragraph.strip())
            pdf.ln(4)
    
    # Save the cover letter
    filename = os.path.join(OUTPUT_DIR, f"{candidate_name.replace(' ', '_')}_Cover_Letter.pdf")
    pdf.output(filename)
    return filename

# Add a route to serve the HTML page
@app.get("/", response_class=HTMLResponse)
async def get_html():
    with open(os.path.join(TEMPLATES_DIR, "index.html"), "r") as f:
        html_content = f.read()
    return HTMLResponse(content=html_content)





# Continuing from the previous code...

@app.post("/upload/resume", response_model=dict)
async def upload_resume(file: UploadFile = File(...)):
    """Upload and process a resume or template file"""
    # Check file extension
    filename = file.filename
    file_extension = filename.split('.')[-1].lower()
    
    if file_extension not in ['pdf', 'docx', 'txt']:
        raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF, DOCX, or TXT files.")
    
    # Save the file
    file_path = os.path.join(OUTPUT_DIR, filename)
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())
    
    # Extract text from the file
    try:
        extracted_text = extract_text_from_file(file_path)
        return {
            "filename": filename,
            "file_path": file_path,
            "status": "success",
            "message": f"File uploaded and processed successfully"
        }
    except Exception as e:
        # If there's an error, remove the uploaded file
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@app.post("/upload/skill-matrix", response_model=dict)
async def upload_skill_matrix(file: UploadFile = File(...)):
    """Upload and process a skill matrix Excel file"""
    # Check file extension
    filename = file.filename
    file_extension = filename.split('.')[-1].lower()
    
    if file_extension not in ['xlsx', 'xls']:
        raise HTTPException(status_code=400, detail="Unsupported file format. Please upload Excel files only.")
    
    # Save the file
    file_path = os.path.join(OUTPUT_DIR, filename)
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())
    
    # Process the Excel file
    try:
        sheet_data = load_skill_matrix(file_path)
        
        # Count total records across all sheets
        total_records = sum(len(sheet["Data"]) for sheet in sheet_data)
        
        return {
            "filename": filename,
            "file_path": file_path,
            "status": "success",
            "sheets_processed": len(sheet_data),
            "total_records": total_records,
            "message": f"Skill matrix uploaded and processed successfully"
        }
    except Exception as e:
        # If there's an error, remove the uploaded file
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Error processing skill matrix: {str(e)}")

@app.post("/search/by-name", response_model=List[dict])
async def search_by_name(query: NameQuery):
    """Search for a person in the skill matrix by name"""
    if not sheets_data:
        raise HTTPException(status_code=400, detail="No skill matrix data loaded. Please upload a skill matrix file first.")
    
    results = get_json_by_name(query.first_name, query.last_name)
    
    # If no results found, try case-insensitive search
    if not results:
        # Try case-insensitive search
        results = []
        for sheet in sheets_data:

            for record in sheet["Data"]:
                # Check both standardized and possible original column names with case insensitivity
                fn_keys = ["First_Name", "First Name"] 
                ln_keys = ["Last_Name", "Last Name"]
                
                fn_match = False
                for key in fn_keys:
                    if key in record and isinstance(record[key], str) and record[key].lower() == query.first_name.lower():
                        fn_match = True
                        break
                
                ln_match = False
                for key in ln_keys:
                    if key in record and isinstance(record[key], str) and record[key].lower() == query.last_name.lower():
                        ln_match = True
                        break
                
                if fn_match and ln_match:
                    # Include sheet name in record for context
                    record_with_sheet = record.copy()
                    record_with_sheet["Sheet Name"] = sheet["Sheet Name"]
                    results.append(record_with_sheet)
    
    return results

@app.post("/generate/resume", response_model=dict)
async def generate_resume(request: ResumeRequest, background_tasks: BackgroundTasks):
    """Generate a new resume based on template and skill matrix data"""
    # Validate that files exist
    if not os.path.exists(request.old_resume_path):
        raise HTTPException(status_code=400, detail="Resume file not found. Please upload a resume first.")
    
    if not os.path.exists(request.template_path):
        raise HTTPException(status_code=400, detail="Template file not found. Please upload a template first.")
    
    try:
        # Extract text from files
        old_resume_text = extract_text_from_file(request.old_resume_path)
        template_text = extract_text_from_file(request.template_path)
        
        # Get skill matrix data if available
        skill_matrix_json = "No skill matrix data available"
        if request.skill_matrix_path and os.path.exists(request.skill_matrix_path):
            # If first name and last name are provided, get specific data
            if request.first_name and request.last_name:
                skill_matrix_data = get_json_by_name(request.first_name, request.last_name)
                if skill_matrix_data:
                    skill_matrix_json = json.dumps(skill_matrix_data)
            else:
                # Log a warning but don't fail
                print("Warning: Skill matrix provided but no name specified")
        
        # Step 1: Extract structured JSON from template
        try:
            template_json_str = LLM_CALL_1(template_text)
            
            # Clean up the JSON string to ensure it's valid
            template_json_str = template_json_str.replace('```json', '').replace('```', '').strip()
            
            try:
                formatted_json_schema = json.loads(template_json_str)
            except json.JSONDecodeError as e:
                # If template JSON parsing fails, use a default schema
                print(f"Error parsing template JSON: {e}, using default schema")
                formatted_json_schema = {
                    "name": "",
                    "designation": "",
                    "objective": "",
                    "education": [],
                    "skills": [],
                    "project_details": {}
                }
        except Exception as e:
            print(f"Error in LLM_CALL_1: {str(e)}")
            # Use default schema if LLM call fails
            formatted_json_schema = {
                "name": "",
                "designation": "",
                "objective": "",
                "education": [],
                "skills": [],
                "project_details": {}
            }
        
        # Step 2: Reformat old resume with skill matrix data
        try:
            new_resume_json = LLM_CALL_2(
                formatted_json_schema, 
                old_resume_text, 
                skill_matrix_json
            )
            
            # Ensure we have valid JSON
            try:
                resume_data = json.loads(new_resume_json)
            except json.JSONDecodeError as e:
                # Clean up the JSON if necessary
                print(f"Error parsing generated JSON: {e}, attempting to clean")
                clean_json = new_resume_json.replace('```json', '').replace('```', '').strip()
                try:
                    resume_data = json.loads(clean_json)
                except json.JSONDecodeError:
                    # If still failing, create minimal valid JSON
                    print("Error parsing cleaned JSON, using minimal data")
                    # Extract name from resume
                    name = "Candidate"
                    for line in old_resume_text.split('\n'):
                        if line.strip() and len(line.strip()) < 50:  # Simple heuristic to find a name
                            name = line.strip()
                            break
                    
                    resume_data = {
                        "name": name,
                        "designation": "Professional",
                        "objective": "Experienced professional seeking new opportunities.",
                        "education": ["Education details not extracted"],
                        "skills": ["Skills not extracted"],
                        "project_details": {
                            "project1": {
                                "name": "Project",
                                "role": "Team Member",
                                "description": "Project description not extracted",
                                "technology": "Technologies not extracted",
                                "role_played": "Role details not extracted"
                            }
                        }
                    }
        except Exception as e:
            print(f"Error in LLM_CALL_2: {str(e)}")
            # Create minimal valid JSON if LLM call fails
            name = "Candidate"
            for line in old_resume_text.split('\n'):
                if line.strip() and len(line.strip()) < 50:
                    name = line.strip()
                    break
                
            resume_data = {
                "name": name,
                "designation": "Professional",
                "objective": "Experienced professional seeking new opportunities.",
                "education": ["Education details not extracted"],
                "skills": ["Skills not extracted"],
                "project_details": {
                    "project1": {
                        "name": "Project",
                        "role": "Team Member",
                        "description": "Project description not extracted",
                        "technology": "Technologies not extracted",
                        "role_played": "Role details not extracted"
                    }
                }
            }
            
        # Step 3: Generate the resume PDF
        try:
            resume_path = generate_resume_1(resume_data)
        except Exception as e:
            print(f"Error generating PDF: {str(e)}")
            # Create a simple PDF with basic information
            pdf = PDF()
            pdf.set_margins(15, 15, 15)
            pdf.add_page()
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 8, resume_data.get("name", "Candidate"), ln=True, align="L")
            pdf.set_font("Arial", "", 11)
            pdf.multi_cell(0, 7, "An error occurred while generating the complete resume.")
            pdf.ln(10)
            pdf.multi_cell(0, 7, "Please try again or contact support if the issue persists.")
            
            resume_path = os.path.join(OUTPUT_DIR, f"{resume_data.get('name', 'Candidate').replace(' ', '_')}_Resume.pdf")
            pdf.output(resume_path)
        
        # Step 4: Generate a cover letter
        cover_letter_status = "Not generated"
        try:
            cover_letter_text = generate_cover_letter_from_resume(resume_data)
            cover_letter_path = generate_cover_letter_pdf(cover_letter_text, resume_data["name"])
            cover_letter_status = f"Generated successfully"
        except Exception as e:
            cover_letter_status = f"Failed to generate"
            print(f"Error generating cover letter: {str(e)}")
        
        return {
            "message": "Resume generated successfully",
            "resume_path": resume_path,
            "cover_letter_status": cover_letter_status
        }
    
    except Exception as e:
        import traceback
        error_detail = str(e)
        print(f"Error in generate_resume endpoint: {error_detail}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error generating resume: {error_detail}")
    


@app.get("/download/{filename}")
async def download_file(filename: str):
    """Endpoint to download generated files"""
    file_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(file_path, filename=filename)

# Create prompts if they don't exist
def create_prompts():
    llm1_path = os.path.join(PROMPTS_DIR, "LLM1.txt")
    if not os.path.exists(llm1_path):
        with open(llm1_path, "w") as f:
            f.write("""
You are an AI assistant that extracts structured information from resumes.
Your task is to analyze the provided resume text and extract key details into a well-structured JSON format.

The JSON structure should include:
{
  "name": "Full Name",
  "designation": "Job Title/Role",
  "objective": "Professional summary/objective statement",
  "education": ["Degree 1", "Degree 2"],
  "skills": ["Skill 1", "Skill 2", "Skill 3"],
  "project_details": {
    "project1": {
      "name": "Project Name",
      "role": "Role in Project",
      "description": "Brief description of the project",
      "technology": "Technologies used",
      "role_played": "Responsibilities and contributions"
    },
    "project2": {
      "name": "Project Name",
      "role": "Role in Project",
      "description": "Brief description of the project",
      "technology": "Technologies used",
      "role_played": "Responsibilities and contributions"
    }
  }
}

Follow these guidelines:
1. Extract information ONLY from the provided text.
2. Do not invent or add details that aren't present in the original text.
3. Format the output as valid JSON that strictly adheres to the structure above.
4. If certain information is missing, use empty strings or arrays rather than omitting fields.
5. Include at least 2-3 projects if available in the resume.
6. Organize skills in a logical manner, grouping similar skills together.
7. Be concise but comprehensive in extracting information.
8. Ensure all JSON is properly formatted and valid.
            """)

# Run startup tasks
@app.on_event("startup")
async def startup_event():
    create_prompts()

# Mount static files directory
# app.mount("/static", StaticFiles(directory=os.path.join(BASE_PATH, "static")), name="static")
static_dir = os.path.join(BASE_PATH, "static")
os.makedirs(static_dir, exist_ok=True)  # Create the directory if it doesn't exist
app.mount("/static", StaticFiles(directory=static_dir), name="static")

# Run the application
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

