


from fastapi import FastAPI, File, UploadFile, Form, HTTPException, BackgroundTasks,Query
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional, Dict, Any
import pandas as pd
import json
import fitz  # PyMuPDF
from docx import Document
import io
import os
from pydantic import BaseModel
from mistralai import Mistral
from fpdf import FPDF


from typing import List

# import json
import re

app = FastAPI(title="Resume Automation API")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# API key for Mistral
api_key = "QsIe8weZaSTfNx9ykWlbjC3vGhWVSzWE"

# Global variable to store loaded skill matrix data
sheets_data = []
names_list = []

# Paths configuration
BASE_PATH = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(BASE_PATH, "Logo.png")
PROMPTS_DIR = os.path.join(BASE_PATH, "Prompts")
OUTPUT_DIR = os.path.join(BASE_PATH, "Output")
TEMPLATES_DIR = os.path.join(BASE_PATH, "templates")

# Create required directories if they don't exist
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(PROMPTS_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)

# Create HTML template directory and file

# Setup templates
templates = Jinja2Templates(directory=TEMPLATES_DIR)

# Pydantic models
class NameQuery(BaseModel):
    first_name: str
    last_name: str




class EmployeeInfo(BaseModel):
    ID: int
    first_name: str
    last_name: str
    sheet_name: str

class ResumeRequest(BaseModel):
    template_path: str
    old_resume_path: Optional[str] = None
    old_cover_letter_path: Optional[str] = None
    cover_letter_path: Optional[str] = None  # Alternative name
    skill_matrix_path: Optional[str] = None
    first_name: Optional[str] = None
    last_name: Optional[str] = None



class EmployeeIdQuery(BaseModel):
    employee_id: int

# Update or make sure the EmployeeInfo model has ID field

    
    # Add a method to get the cover letter path from either field
    def get_cover_letter_path(self):
        return self.old_cover_letter_path or self.cover_letter_path


# File extraction functions
def extract_text_from_file(file_path: str) -> str:
    """Extract text from various file formats using file path"""
    file_extension = file_path.split('.')[-1].lower()
    text = ''
    
    if file_extension == 'pdf':
        # Extract text from PDF using PyMuPDF
        with open(file_path, 'rb') as file:
            doc = fitz.open(stream=file.read(), filetype='pdf')
            for page in doc:
                text += page.get_text("text")
    elif file_extension == 'docx':
        # Extract text from DOCX
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    elif file_extension == 'txt':
        # Extract text from TXT
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    else:
        raise ValueError('Unsupported file format. Please provide PDF, DOCX, or TXT files.')
    
    return text

def extract_text_from_bytes(file_content: bytes, file_extension: str) -> str:
    """Extract text from various file formats using file bytes"""
    text = ''
    
    if file_extension == 'pdf':
        # Extract text from PDF using PyMuPDF
        doc = fitz.open(stream=file_content, filetype='pdf')
        for page in doc:
            text += page.get_text("text")
    elif file_extension == 'docx':
        # Extract text from DOCX
        doc = Document(io.BytesIO(file_content))
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    elif file_extension == 'txt':
        # Extract text from TXT
        text = file_content.decode('utf-8')
    else:
        raise ValueError('Unsupported file format. Please provide PDF, DOCX, or TXT files.')
    
    return text






# Excel processing functions
def load_skill_matrix(file_path: str) -> List[dict]:
    """Load and process skill matrix Excel file from path"""
    global sheets_data
    sheets_data = []
    
    # Load Excel file
    xls = pd.ExcelFile(file_path)
    record_id = 1  # Unique number for each entry
    
    # Process all sheets from the second one onward
    for sheet_name in xls.sheet_names[1:]:  # Skip the first sheet
        df = pd.read_excel(xls, sheet_name=sheet_name)
        
        # Standardize column names to ensure consistency
        if len(df.columns) >= 4:
            # Create standardized column names
            std_columns = ['First_Name', 'Last_Name', 'Experience', 'Expertise'] + list(df.columns[4:])
            
            # Map original columns to standardized names
            column_mapping = {original: standard for original, standard in zip(df.columns, std_columns)}
            df = df.rename(columns=column_mapping)
        
        # Convert each sheet's data to a list of dictionaries with unique numbers
        records = []
        for _, row in df.iterrows():
            record = {"ID": record_id}  # Assign a unique ID
            record.update(row.to_dict())  # Convert row data to dictionary
            records.append(record)
            record_id += 1  # Increment unique ID
        
        sheet_dict = {
            "Sheet Name": sheet_name,
            "Data": records
        }
        
        sheets_data.append(sheet_dict)
    
    return sheets_data

def load_skill_matrix_from_bytes(file_content: bytes) -> List[dict]:
    """Load and process skill matrix Excel file from bytes"""
    global sheets_data
    sheets_data = []
    
    # Load Excel file from bytes
    xls = pd.ExcelFile(io.BytesIO(file_content))
    record_id = 1  # Unique number for each entry
    
    # Process all sheets from the second one onward
    for sheet_name in xls.sheet_names[1:]:  # Skip the first sheet
        df = pd.read_excel(xls, sheet_name=sheet_name)
        
        # Standardize column names to ensure consistency
        if len(df.columns) >= 4:
            # Create standardized column names
            std_columns = ['First_Name', 'Last_Name', 'Experience', 'Expertise'] + list(df.columns[4:])
            
            # Map original columns to standardized names
            column_mapping = {original: standard for original, standard in zip(df.columns, std_columns)}
            df = df.rename(columns=column_mapping)
        
        # Convert each sheet's data to a list of dictionaries with unique numbers
        records = []
        for _, row in df.iterrows():
            record = {"ID": record_id}  # Assign a unique ID
            record.update(row.to_dict())  # Convert row data to dictionary
            records.append(record)
            record_id += 1  # Increment unique ID
        
        sheet_dict = {
            "Sheet Name": sheet_name,
            "Data": records
        }
        
        sheets_data.append(sheet_dict)
    
    return sheets_data





def get_json_by_name(first_name: str, last_name: str) -> List[dict]:
    """Search for records by first and last name"""
    results = []
    for sheet in sheets_data:
        for record in sheet["Data"]:
            # Check both standardized and possible original column names with case insensitivity
            fn_keys = ["First_Name", "First Name"] 
            ln_keys = ["Last_Name", "Last Name"]
            
            fn_match = False
            for key in fn_keys:
                if key in record and isinstance(record[key], str) and record[key].lower() == first_name.lower():
                    fn_match = True
                    break
            
            ln_match = False
            for key in ln_keys:
                if key in record and isinstance(record[key], str) and record[key].lower() == last_name.lower():
                    ln_match = True
                    break
            
            if fn_match and ln_match:
                # Include sheet name in record for context
                record_with_sheet = record.copy()
                record_with_sheet["Sheet Name"] = sheet["Sheet Name"]
                results.append(record_with_sheet)
    
    return results



# LLM functions
def LLM_CALL_1(resume_text: str) -> str:
    """Extract structured JSON from resume text"""
    try:
        model = "mistral-large-latest"
        client = Mistral(api_key=api_key)
        parsed_text = resume_text[:5000]  # Limit text length to avoid token limits
        
        # Read system prompt from file
        system_prompt_path = os.path.join(PROMPTS_DIR, "LLM1.txt")
        if os.path.exists(system_prompt_path):
            with open(system_prompt_path, 'r') as file:
                system_prompt = file.read()
        else:
            # Fallback if file doesn't exist
            system_prompt = """
            You are an AI assistant that extracts structured information from resumes.
            Extract relevant details from the resume text and format them into a well-structured JSON object.
            Include name, designation, objective, education, skills, and project details.
            Your response should be a valid JSON object without any additional text or markdown formatting.
            """
        
        try:
            chat_response = client.chat.complete(
                model=model,
                temperature=0,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Here is an extracted resume text:\n\n{parsed_text}\n\nFormat it into JSON."}
                ],
                response_format = {
                    "type": "json_object",
                }
            )
            
            json_schema = chat_response.choices[0].message.content
            
            # Clean the JSON output by removing triple backticks if present
            cleaned_json = json_schema.strip()
            if cleaned_json.startswith("```json") and cleaned_json.endswith("```"):
                cleaned_json = cleaned_json[7:-3].strip()
            elif cleaned_json.startswith("```") and cleaned_json.endswith("```"):
                cleaned_json = cleaned_json[3:-3].strip()
            
            # Validate the JSON
            try:
                json.loads(cleaned_json)
                return cleaned_json
            except json.JSONDecodeError:
                # If invalid JSON, try to fix common issues
                import re
                cleaned_json = cleaned_json.replace('\n', ' ').replace('\r', '')
                cleaned_json = re.sub(r'(?<!")(\w+)(?=":)', r'"\1"', cleaned_json)
                return cleaned_json
                
        except Exception as e:
            print(f"Error in Mistral API call: {str(e)}")
            # Return a default JSON structure
            default_json = {
                "name": "",
                "designation": "",
                "objective": "",
                "education": [],
                "skills": [],
                "project_details": {}
            }
            return json.dumps(default_json)
    
    except Exception as e:
        print(f"Unexpected error in LLM_CALL_1: {str(e)}")
        default_json = {
            "name": "",
            "designation": "",
            "objective": "",
            "education": [],
            "skills": [],
            "project_details": {}
        }
        return json.dumps(default_json)
    

    # old_resume + skill matrix +


def LLM_CALL_2(formatted_json_schema: Dict, old_resume_text: str, skill_matrix_json: str) -> str:
    """Reformat old resume with skill matrix data into new structured JSON"""
    try:
        model = "mistral-large-latest"
        client = Mistral(api_key=api_key)
        
        # Limit text lengths to avoid token limits
        old_resume_text = old_resume_text[:4000]
        if len(skill_matrix_json) > 2000:
            skill_matrix_json = skill_matrix_json[:2000] + "... [truncated]"
        
        system_prompt = """
        You are an AI assistant that reformats resumes into a structured JSON format.
        Take the competency matrix and old resume as input, extract relevant details, and map them to the new resume format.
        Only respond with the new resume format as a JSON object that adheres strictly to the provided JSON Schema.
        Your response must be a valid, properly formatted JSON object without any additional text, explanations, or markdown formatting.
        """
        
        try:
            chat_response = client.chat.complete(
                model=model,
                temperature=0,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"""
                    You are given the following details:
                    
                    **Skill Matrix (Extracted from External Data Sources):**
                    {skill_matrix_json}
                    
                    **Old Resume (Extracted Text):**
                    {old_resume_text}
                    
                    Your task is to extract relevant details and format them into the following structured JSON format:
                    
                    {json.dumps(formatted_json_schema, indent=2)}
                    
                    **Instructions:**
                    - Only use the given details; do not generate fictional information.
                    - Follow the JSON structure exactly as provided.
                    - Ensure the output is a valid JSON object (no extra text or explanations).
                    - "name" and "designation" are mandatory fields.
                    - Create a proper summarized objective.
                    
                    Generate the updated resume in JSON format:
                    """}
                ],
                response_format = {
                    "type": "json_object",
                }
            )
            
            resume_content = chat_response.choices[0].message.content
            
            # Attempt to validate and clean JSON if needed
            try:
                json.loads(resume_content)
                return resume_content
            except json.JSONDecodeError:
                # Clean up if it's not valid JSON
                cleaned_content = resume_content.replace('```json', '').replace('```', '').strip()
                return cleaned_content
                
        except Exception as e:
            print(f"Error in Mistral API call: {str(e)}")
            # Extract name from resume text as fallback
            name = "Candidate"
            for line in old_resume_text.split('\n'):
                if line.strip() and len(line.strip()) < 50:
                    name = line.strip()
                    break
                    
            # Return a basic JSON with resume text information
            default_json = {
                "name": name,
                "designation": "Professional",
                "objective": "Experienced professional seeking new opportunities.",
                "education": ["Education details not extracted"],
                "skills": ["Skills not extracted"],
                "project_details": {
                    "project1": {
                        "name": "Project",
                        "role": "Team Member",
                        "description": "Project description not extracted",
                        "technology": "Technologies not extracted",
                        "role_played": "Role details not extracted"
                    }
                }
            }
            return json.dumps(default_json)
    
    except Exception as e:
        print(f"Unexpected error in LLM_CALL_2: {str(e)}")
        # Return a basic JSON structure
        default_json = {
            "name": "Candidate",
            "designation": "Professional",
            "objective": "Experienced professional seeking new opportunities.",
            "education": ["Education details not extracted"],
            "skills": ["Skills not extracted"],
            "project_details": {
                "project1": {
                    "name": "Project",
                    "role": "Team Member",
                    "description": "Project description not extracted",
                    "technology": "Technologies not extracted",
                    "role_played": "Role details not extracted"
                }
            }
        }
        return json.dumps(default_json)



# old_resume + skill matrix
# api_key = "QsIe8weZaSTfNx9ykWlbjC3vGhWVSzWE"

# def LLM_CALL_2(formatted_json_schema, old_resume_text, json_output):
   
#     model = "mistral-large-latest"

#     client = Mistral(api_key=api_key)

#     new_resume_text = formatted_json_schema
#     Content = old_resume_text
#     Skill_Matrix = json_output
    

#     """
#     Calls an LLM API (via OpenRouter using ASUS TUF API) to reformat the old resume plus candidate skill data
#     into a new structured JSON following the provided new resume format.
#     """
#     system_prompt = f"""
#     You are an AI assistant that reformats resumes into a structured JSON format.
#     Take the competency matrix and old resume as input, extract relevant details, and map them to the new resume format.
#     Only respond with the new resume format as a JSON object that adheres strictly to the provided JSON Schema. Do not include any extra messages.
#      Your task is to extract relevant details and format them into the following structured JSON format:
            
#             {new_resume_text}

#             **Instructions:**
#             - Only use the given details; do not generate fictional information.
#             - Follow the JSON structure exactly as provided.
#             - Ensure the output starts directly with a valid JSON object (no extra text or explanations).
            
#             Generate the updated resume in JSON format:
#             name and designation mandatory key value pairs and create a proper summarized objective a key  in for any kind of templete.
#             Points to note : In the key "Education" form a sentance and use that instead of having multiple sub sections. Do not make up values , if you have no access to a value, don't make it up.
#             You are given the following details: 
           
#     """
    

#     chat_response = client.chat.complete(
#         model=model,
#         temperature=0.2,
#         messages=[
#             {"role": "system", "content": system_prompt},
#             {"role": "user", "content": f"""
            
#             **Skill Matrix (Extracted from External Data Sources):**
#             {Skill_Matrix}

#             **Old Resume (Extracted Text):**
#             {Content}

#             """}
#         ],
#         response_format = {
#           "type": "json_object",
#       }
#     )


#     resume_content = chat_response.choices[0].message.content

#     return resume_content





# skill_matrix+cover_letter


def LLM_CALL_3(formatted_json_schema, old_cover_text, json_output):
   
    model = "mistral-large-latest"

    client = Mistral(api_key=api_key)

    new_resume_text = formatted_json_schema
    Content2 = old_cover_text
    Skill_Matrix = json_output
    

    """
    Calls an LLM API (via OpenRouter using ASUS TUF API) to reformat the old resume plus candidate skill data
    into a new structured JSON following the provided new resume format.
    """
    system_prompt = f"""
    You are an AI assistant that reformats resumes into a structured JSON format.
    Take the competency matrix and old resume as input, extract relevant details, and map them to the new resume format.
    Only respond with the new resume format as a JSON object that adheres strictly to the provided JSON Schema. Do not include any extra messages.
     Your task is to extract relevant details and format them into the following structured JSON format:
            
            {new_resume_text}

            **Instructions:**
            - Only use the given details; do not generate fictional information.
            - Follow the JSON structure exactly as provided.
            - Ensure the output starts directly with a valid JSON object (no extra text or explanations).
            
            Generate the updated resume in JSON format:
            name and designation mandatory key value pairs and create a proper summarized objective a key  in for any kind of templete.
            Points to note : In the key "Education" form a sentance and use that instead of having multiple sub sections. Do not make up values , if you have no access to a value, don't make it up.
            You are given the following details: 
           
    """
    

    chat_response = client.chat.complete(
        model=model,
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"""
            
            **Skill Matrix (Extracted from External Data Sources):**
            {Skill_Matrix}

            **Old Coverletter (Extracted Text):**
            {Content2}

            """}
        ],
        response_format = {
          "type": "json_object",
      }
    )


    resume_content = chat_response.choices[0].message.content

    return resume_content



# old_resume+old_coverletter

def LLM_CALL_4(formatted_json_schema, old_resume_text, old_cover_text):
   
    model = "mistral-large-latest"

    client = Mistral(api_key=api_key)

    new_resume_text = formatted_json_schema
    Content = old_resume_text
    Content2 = old_cover_text

    """
    Calls an LLM API (via OpenRouter using ASUS TUF API) to reformat the old resume plus candidate skill data
    into a new structured JSON following the provided new resume format.
    """
    system_prompt = f"""
    You are an AI assistant that reformats resumes into a structured JSON format.
    Take the competency matrix and old resume as input, extract relevant details, and map them to the new resume format.
    Only respond with the new resume format as a JSON object that adheres strictly to the provided JSON Schema. Do not include any extra messages.
     Your task is to extract relevant details and format them into the following structured JSON format:
            
            {new_resume_text}

            **Instructions:**
            - Only use the given details; do not generate fictional information.
            - Follow the JSON structure exactly as provided.
            - Ensure the output starts directly with a valid JSON object (no extra text or explanations).
            
            Generate the updated resume in JSON format:
            name and designation mandatory key value pairs and create a proper summarized objective a key  in for any kind of templete.
            Points to note : In the key "Education" form a sentance and use that instead of having multiple sub sections. Do not make up values , if you have no access to a value, don't make it up.
            You are given the following details: 
           
    """
    

    chat_response = client.chat.complete(
        model=model,
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"""
            
            
            **Old Resume (Extracted Text):**
            {Content}


            **Old Coverletter (Extracted Text):**
            {Content2}

            """}
        ],
        response_format = {
          "type": "json_object",
      }
    )


    resume_content = chat_response.choices[0].message.content

    return resume_content





# old resume

def LLM_CALL_5(formatted_json_schema, old_resume_text):
   
    model = "mistral-large-latest"

    client = Mistral(api_key=api_key)

    new_resume_text = formatted_json_schema
    Content = old_resume_text
  


    """
    Calls an LLM API (via OpenRouter using ASUS TUF API) to reformat the old resume plus candidate skill data
    into a new structured JSON following the provided new resume format.
    """
    system_prompt = f"""
    You are an AI assistant that reformats resumes into a structured JSON format.
    Take the competency matrix and old resume as input, extract relevant details, and map them to the new resume format.
    Only respond with the new resume format as a JSON object that adheres strictly to the provided JSON Schema. Do not include any extra messages.
    **Instructions:**
            - Only use the given details; do not generate fictional information.
            - Follow the JSON structure exactly as provided.
            - Ensure the output starts directly with a valid JSON object (no extra text or explanations).
            - Try to keep the project description within 40-50 words.

            Generate the updated resume in JSON format:
            name and designation mandatory key value pairs and create a proper summarized objective a key  in for any kind of templete.
            make sure it follows the given format of formatted_json_schema wiht same json format with key value pairs. 
            Points to note : In the key "Education" form a sentance and use that instead of having multiple sub sections. Do not make up values , if you have no access to a value, don't make it up.
            Keep the names of the keyvalues in the json schema same.Keep the key for title of project  as name only. Skills also give as per the json schema if no sub sections give as list.

            You are given the following details: You are given the following details:

            Your task is to extract relevant details and format them into the following structured JSON format:
            {new_resume_text}
    """
    

    chat_response = client.chat.complete(
        model=model,
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"""
             
            **Old Resume (Extracted Text):**
            {Content}
           

            """}
        ],
        response_format = {
          "type": "json_object",
      }
    )


    resume_content = chat_response.choices[0].message.content

    return resume_content
# 


# old cover letter





def LLM_CALL_6(formatted_json_schema, old_cover_text):
   
    model = "mistral-large-latest"

    client = Mistral(api_key=api_key)

    new_resume_text = formatted_json_schema
    Content2 = old_cover_text
  

    """
    Calls an LLM API (via OpenRouter using ASUS TUF API) to reformat the old resume plus candidate skill data
    into a new structured JSON following the provided new resume format.
    """
    system_prompt = f"""
    You are an AI assistant that reformats resumes into a structured JSON format.
    Take the competency matrix and old resume as input, extract relevant details, and map them to the new resume format.
    Only respond with the new resume format as a JSON object that adheres strictly to the provided JSON Schema. Do not include any extra messages.
    **Instructions:**
            - Only use the given details; do not generate fictional information.
            - Follow the JSON structure exactly as provided.
            - Ensure the output starts directly with a valid JSON object (no extra text or explanations).

            Generate the updated resume in JSON format:
            name and designation mandatory key value pairs and create a proper summarized objective a key  in for any kind of templete.
            make sure it follows the given format of formatted_json_schema wiht same json format with key value pairs. 
            Points to note : In the key "Education" form a sentance and use that instead of having multiple sub sections. Do not make up values , if you have no access to a value, don't make it up.
            You are given the following details: You are given the following details:

            Your task is to extract relevant details and format them into the following structured JSON format:
            {new_resume_text}
    """
    

    chat_response = client.chat.complete(
        model=model,
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"""
             
            **Old CoverLetter (Extracted Text):**
            {Content2}
           

            """}
        ],
        response_format = {
          "type": "json_object",
      }
    )


    resume_content = chat_response.choices[0].message.content

    return resume_content




# only skill matrix

def LLM_CALL_7(formatted_json_schema, json_output):
   
    model = "mistral-large-latest"

    client = Mistral(api_key=api_key)

    new_resume_text = formatted_json_schema
    Skill_Matrix = json_output
    

    """
    Calls an LLM API (via OpenRouter using ASUS TUF API) to reformat the old resume plus candidate skill data
    into a new structured JSON following the provided new resume format.
    """
    system_prompt = f"""
    You are an AI assistant that reformats resumes into a structured JSON format.
    Take the competency matrix and old resume as input, extract relevant details, and map them to the new resume format.
    Only respond with the new resume format as a JSON object that adheres strictly to the provided JSON Schema. Do not include any extra messages.
     Your task is to extract relevant details and format them into the following structured JSON format:
            
            {new_resume_text}

            **Instructions:**
            - Only use the given details; do not generate fictional information.
            - Follow the JSON structure exactly as provided.
            - Ensure the output starts directly with a valid JSON object (no extra text or explanations).
            
            Generate the updated resume in JSON format:
            name and designation mandatory key value pairs and create a proper summarized objective a key  in for any kind of templete.
            Points to note : In the key "Education" form a sentance and use that instead of having multiple sub sections. Do not make up values , if you have no access to a value, don't make it up.
            You are given the following details: 
           
    """
    

    chat_response = client.chat.complete(
        model=model,
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"""
            
            **Skill Matrix (Extracted from External Data Sources):**
            {Skill_Matrix}

            

            """}
        ],
        response_format = {
          "type": "json_object",
      }
    )


    resume_content = chat_response.choices[0].message.content

    return resume_content


#   skill matrix+old resume+old cover letter




def LLM_CALL_8(formatted_json_schema, old_resume_text, old_cover_text, json_output):
   
    model = "mistral-large-latest"

    client = Mistral(api_key=api_key)

    new_resume_text = formatted_json_schema
    Content = old_resume_text
    Content2 = old_cover_text
    Skill_Matrix = json_output
    

    """
    Calls an LLM API (via OpenRouter using ASUS TUF API) to reformat the old resume plus candidate skill data
    into a new structured JSON following the provided new resume format.
    """
    system_prompt = f"""
    You are an AI assistant that reformats resumes into a structured JSON format.
    Take the competency matrix and old resume as input, extract relevant details, and map them to the new resume format.
    Only respond with the new resume format as a JSON object that adheres strictly to the provided JSON Schema. Do not include any extra messages.
     Your task is to extract relevant details and format them into the following structured JSON format:
            
            {new_resume_text}

            **Instructions:**
            - Only use the given details; do not generate fictional information.
            - Follow the JSON structure exactly as provided.
            - Ensure the output starts directly with a valid JSON object (no extra text or explanations).
            
            Generate the updated resume in JSON format:
            name and designation mandatory key value pairs and create a proper summarized objective a key  in for any kind of templete.
            Points to note : In the key "Education" form a sentance and use that instead of having multiple sub sections. Do not make up values , if you have no access to a value, don't make it up.
            You are given the following details: 
           
    """
    

    chat_response = client.chat.complete(
        model=model,
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"""
            
            
            **Skill Matrix (Extracted from External Data Sources):**
            {Skill_Matrix}

            **Old Resume (Extracted Text):**
            {Content}


            **Old Coverletter (Extracted Text):**
            {Content2}

            """}
        ],
        response_format = {
          "type": "json_object",
      }
    )


    resume_content = chat_response.choices[0].message.content

    return resume_content



def generate_cover_letter_from_resume(resume_summary: Dict) -> str:
    """Generates a professional cover letter based on the candidate's resume summary"""
    model = "mistral-large-latest"
    client = Mistral(api_key=api_key)
    
    system_prompt = """
    You are an AI assistant that creates professional cover letters based only on a candidate's resume summary.
    Your task is to analyze the resume summary, infer the candidate's expertise, and generate a well-structured cover letter.
    Ensure the tone is formal, engaging, and professional.
    """
    
    user_prompt = f"""
    Candidate Resume Summary:
    {json.dumps(resume_summary, indent=2)}
    
    Based on this summary, write a professional cover letter.
    Follow this structure:
    - Address the hiring manager (use "Dear Hiring Manager" if no specific name is provided).
    - Introduce the candidate and express general interest in roles that match their expertise.
    - Highlight key experiences, achievements, and skills relevant to their domain.
    - Conclude with enthusiasm and a call to action.
    
    Ensure the letter is formal and engaging.
    """
    
    chat_response = client.chat.complete(
        model=model,
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        response_format={"type": "text"}
    )
    
    cover_letter = chat_response.choices[0].message.content
    return cover_letter






class PDF(FPDF):
    def header(self):
        """Adds a logo if available."""
        try:
            self.image(r"C:\Users\CVHS\Downloads\NEW_RESUMATE\server\Logo.png", 160, 8, 35)
        except:
            pass  # Skip if logo is not found
        self.set_font("Arial", "B", 12)
        self.ln(10)

    def chapter_title(self, title):
        self.set_text_color(0, 0, 180)
        self.set_font("Arial", "B", 11)
        self.cell(0, 8, self.clean_text(title), ln=True, align="L")
        self.ln(4)
        self.set_text_color(0, 0, 0)

    def add_project_table(self, project):
        """Creates a vertical-style project details table."""
        col_widths = [50, 120]
        row_height = 8

        self.set_font("Arial", "B", 9)
        self.cell(0, 6, f"<{self.clean_text(project.get('name', 'N/A'))}>", ln=True)
        self.ln(2)
        self.set_font("Arial", "", 9)

        def add_row(label, text):
            if text:
                text = self.clean_text(text)
                text_height = (self.get_string_width(text) // col_widths[1] + 1) * row_height
                self.cell(col_widths[0], text_height, label, border=1, align="C")
                self.multi_cell(col_widths[1], row_height, text, border=1)

        fields = [
            ("Project Name", project.get("name", "")),
            ("Role", project.get("role", "")),
            ("Description", project.get("description", "")),
            ("Technology", project.get("technology", "")),
            ("Role Played", project.get("role_played", "")),
        ]

        for label, text in fields:
            add_row(label, text)

        self.ln()


    def cell(self, w, h=0, txt='', border=0, ln=0, align='', fill=False, link=''):
        txt = self.clean_text(txt)
        super().cell(w, h, txt, border, ln, align, fill, link)

    def multi_cell(self, w, h, txt, border=0, align='J', fill=False):
        txt = self.clean_text(txt)
        super().multi_cell(w, h, txt, border, align, fill)

    def clean_text(self, text):
        if not isinstance(text, str):
            text = str(text)
        text = re.sub(r"[“”]", '"', text)
        text = re.sub(r"[‘’]", "'", text)
        return text.encode("latin-1", "replace").decode("latin-1")



def generate_resume_1(data_json: str) -> str:

    """Generates the resume PDF with precise formatting."""
    # Handle both string and dictionary input types
    if isinstance(data_json, str):
        data = json.loads(data_json)
    else:
        data = data_json
    

    pdf = PDF()
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    pdf.logo_path = data.get(r"C:\Users\CVHS\Downloads\NEW_RESUMATE\server\Logo.png","")




    if "name" in data and data["name"]:
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 8, pdf.clean_text(data["name"]), ln=True, align="L")
    
    if "designation" in data and data["designation"]:
        pdf.set_font("Arial", "I", 12)
        pdf.cell(0, 8, pdf.clean_text(data["designation"]), ln=True, align="L")
    
    pdf.set_line_width(0.5)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(8)
    
    if "summarized_objective" in data and data["summarized_objective"]:
        pdf.chapter_title("PROFESSIONAL SUMMARY")
        pdf.set_font("Arial", "", 9)
        pdf.multi_cell(0, 7, pdf.clean_text(data["summarized_objective"]))
        pdf.ln(5)
    
    if "education" in data and data["education"]:
        pdf.chapter_title("EDUCATION")
        pdf.set_font("Arial", "", 9)
        pdf.multi_cell(0, 6, pdf.clean_text("".join(data["education"])))
        pdf.ln(5)
    
    if "skills" in data and data["skills"]:
        pdf.chapter_title("SKILLS")
        pdf.set_font("Arial", "", 9)
        pdf.multi_cell(0, 6, pdf.clean_text(",".join(data["skills"])))
        pdf.ln(8)
        
    if "project_details" in data and data["project_details"]:
        pdf.chapter_title("PROJECT DETAILS")

        # Handle both list of dictionaries and dictionary of dictionaries
        if isinstance(data["project_details"], list):
            projects = data["project_details"]
        elif isinstance(data["project_details"], dict):
            projects = data["project_details"].values()  # Extract values from dict

        for project in projects:
            if isinstance(project, dict):  # Ensure each item is a dictionary
                pdf.add_project_table(project)

    filename = os.path.join(OUTPUT_DIR, f"{pdf.clean_text(data.get('name', 'Resume').replace(' ', '_'))}_Resume.pdf")
    pdf.output(filename)
    return filename



def generate_cover_letter_pdf(cover_letter_text: str, candidate_name: str) -> str:
    """Generates a PDF file from the cover letter text."""
    pdf = PDF()
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    
    # Header with candidate name
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 8, f"{candidate_name} - Cover Letter", ln=True, align="L")
    
    # Horizontal Line
    pdf.set_line_width(0.5)
    pdf.line(15, pdf.get_y(), 195, pdf.get_y())
    pdf.ln(8)
    
    # Cover letter content
    pdf.set_font("Arial", "", 11)
    
    # Split text into paragraphs and add to PDF
    paragraphs = cover_letter_text.split("\n\n")
    for paragraph in paragraphs:
        if paragraph.strip():
            pdf.multi_cell(0, 7, paragraph.strip())
            pdf.ln(4)
    
    # Save the cover letter
    filename = os.path.join(OUTPUT_DIR, f"{candidate_name.replace(' ', '_')}_Cover_Letter.pdf")
    pdf.output(filename)
    return filename

# Add a route to serve the HTML page
@app.get("/", response_class=HTMLResponse)
async def get_html():
    with open(os.path.join(TEMPLATES_DIR, "index.html"), "r") as f:
        html_content = f.read()
    return HTMLResponse(content=html_content)





def get_json_by_id(employee_id: int) -> List[dict]:
    """Search for records by employee ID"""
    results = []
    for sheet in sheets_data:
        for record in sheet["Data"]:
            if "ID" in record and record["ID"] == employee_id:
                # Include sheet name in record for context
                record_with_sheet = record.copy()
                record_with_sheet["Sheet Name"] = sheet["Sheet Name"]
                results.append(record_with_sheet)
                # Since ID should be unique, we can break after finding the match
                return results
    
    return results





# Continuing from the previous code...




@app.post("/extract-employees", response_model=dict)
async def extract_employees(skillMatrix: UploadFile = File(...)):
    """Extract employees from uploaded skill matrix file"""
    # Check file extension
    filename = skillMatrix.filename
    file_extension = filename.split('.')[-1].lower()
    
    if file_extension not in ['xlsx', 'xls']:
        raise HTTPException(status_code=400, detail="Unsupported file format. Please upload Excel files only.")
    
    # Read the file content
    content = await skillMatrix.read()
    
    try:
        # Process the Excel file from bytes
        sheets_data = load_skill_matrix_from_bytes(content)
        
        # Extract employee list
        employees = []
        
        for sheet in sheets_data:
            sheet_name = sheet["Sheet Name"]
            for record in sheet["Data"]:
                # Extract first name and last name considering both standardized and original column names
                first_name = None
                for key in ["First_Name", "First Name"]:
                    if key in record and isinstance(record[key], str):
                        first_name = record[key]
                        break
                
                last_name = None
                for key in ["Last_Name", "Last Name"]:
                    if key in record and isinstance(record[key], str):
                        last_name = record[key]
                        break
                
                if first_name and last_name:
                    employees.append({
                        "ID": record["ID"],
                        "first_name": first_name,
                        "last_name": last_name,
                        "sheet_name": sheet_name,
                        "full_name": f"{first_name} {last_name}"
                    })
        
        return {
            "success": True,
            "employees": employees,
            "message": f"Successfully extracted {len(employees)} employees from skill matrix"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error extracting employees: {str(e)}")

@app.post("/upload/resume", response_model=dict)
async def upload_resume(file: UploadFile = File(...)):
    """Upload and process a resume or template file"""
    # Check file extension
    filename = file.filename
    file_extension = filename.split('.')[-1].lower()
    
    if file_extension not in ['pdf', 'docx', 'txt']:
        raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF, DOCX, or TXT files.")
    
    # Save the file
    file_path = os.path.join(OUTPUT_DIR, filename)
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())
    
    # Extract text from the file
    try:
        extracted_text = extract_text_from_file(file_path)
        return {
            "filename": filename,
            "file_path": file_path,
            "status": "success",
            "message": f"File uploaded and processed successfully"
        }
    except Exception as e:
        # If there's an error, remove the uploaded file
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@app.post("/upload/skill-matrix", response_model=dict)
async def upload_skill_matrix(file: UploadFile = File(...)):
    """Upload and process a skill matrix Excel file"""
    # Check file extension
    filename = file.filename
    file_extension = filename.split('.')[-1].lower()
    
    if file_extension not in ['xlsx', 'xls']:
        raise HTTPException(status_code=400, detail="Unsupported file format. Please upload Excel files only.")
    
    # Save the file
    file_path = os.path.join(OUTPUT_DIR, filename)
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())
    
    # Process the Excel file
    try:
        sheet_data = load_skill_matrix(file_path)
        
        # Count total records across all sheets
        total_records = sum(len(sheet["Data"]) for sheet in sheet_data)
        
        return {
            "filename": filename,
            "file_path": file_path,
            "status": "success",
            "sheets_processed": len(sheet_data),
            "total_records": total_records,
            "message": f"Skill matrix uploaded and processed successfully"
        }
    except Exception as e:
        # If there's an error, remove the uploaded file
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Error processing skill matrix: {str(e)}")
    





@app.get("/employees/list", response_model=List[EmployeeInfo])
async def get_employee_list():
    """Get the list of all employees from the loaded skill matrix"""
    if not sheets_data:
        raise HTTPException(status_code=400, detail="No skill matrix data loaded. Please upload a skill matrix file first.")
    
    employee_list = []
    
    for sheet in sheets_data:
        sheet_name = sheet["Sheet Name"]
        for record in sheet["Data"]:
            # Extract first name and last name considering both standardized and original column names
            first_name = None
            for key in ["First_Name", "First Name"]:
                if key in record and isinstance(record[key], str):
                    first_name = record[key]
                    break
            
            last_name = None
            for key in ["Last_Name", "Last Name"]:
                if key in record and isinstance(record[key], str):
                    last_name = record[key]
                    break
            
            if first_name and last_name:
                employee_list.append(
                    EmployeeInfo(
                        ID=record["ID"],
                        first_name=first_name,
                        last_name=last_name,
                        sheet_name=sheet_name
                    )
                )
    
    return employee_list



@app.post("/search/by-id", response_model=List[dict])
async def search_by_id(query: EmployeeIdQuery):
    """Search for a person in the skill matrix by ID"""
    if not sheets_data:
        raise HTTPException(status_code=400, detail="No skill matrix data loaded. Please upload a skill matrix file first.")
    
    results = get_json_by_id(query.employee_id)
    
    if not results:
        raise HTTPException(status_code=404, detail=f"No employee found with ID {query.employee_id}")
    
    return results

@app.post("/search/by-name", response_model=List[dict])
async def search_by_name(query: NameQuery):
    """Search for a person in the skill matrix by name"""
    if not sheets_data:
        raise HTTPException(status_code=400, detail="No skill matrix data loaded. Please upload a skill matrix file first.")
    
    results = get_json_by_name(query.first_name, query.last_name)
    
    # If no results found, try case-insensitive search
    if not results:
        # Try case-insensitive search
        results = []
        for sheet in sheets_data:

            for record in sheet["Data"]:
                # Check both standardized and possible original column names with case insensitivity
                fn_keys = ["First_Name", "First Name"] 
                ln_keys = ["Last_Name", "Last Name"]
                
                fn_match = False
                for key in fn_keys:
                    if key in record and isinstance(record[key], str) and record[key].lower() == query.first_name.lower():
                        fn_match = True
                        break
                
                ln_match = False
                for key in ln_keys:
                    if key in record and isinstance(record[key], str) and record[key].lower() == query.last_name.lower():
                        ln_match = True
                        break
                
                if fn_match and ln_match:
                    # Include sheet name in record for context
                    record_with_sheet = record.copy()
                    record_with_sheet["Sheet Name"] = sheet["Sheet Name"]
                    results.append(record_with_sheet)
    
    return results






# @app.post("/upload/skill-matrix", response_model=dict)
# async def upload_skill_matrix(file: UploadFile = File(...), sheets: Optional[List[str]] = Query(None)):
#     filename = file.filename
#     file_extension = filename.split('.')[-1].lower()

#     if file_extension not in ['xlsx', 'xls']:
#         raise HTTPException(status_code=400, detail="Unsupported file format. Please upload Excel files only.")

#     file_path = os.path.join(OUTPUT_DIR, filename)
#     with open(file_path, "wb") as buffer:
#         buffer.write(await file.read())

#     try:
#         global skill_matrix_data
#         skill_matrix_data = load_skill_matrix(file_path, selected_sheets=sheets)

#         return {
#             "filename": filename,
#             "file_path": file_path,
#             "status": "success",
#             "total_names": len(skill_matrix_data),
#             "message": "Skill matrix uploaded and names extracted successfully"
#         }
#     except Exception as e:
#         os.remove(file_path)
#         raise HTTPException(status_code=500, detail=f"Error processing skill matrix: {str(e)}")


# @app.get("/upload/names", response_model=List[dict])
# async def get_names():
#     if not names_list:
#         raise HTTPException(status_code=400, detail="No names found. Please upload a skill matrix file first.")
#     return names_list



# @app.post("/generate/resume", response_model=dict)
# async def generate_resume(request: ResumeRequest, background_tasks: BackgroundTasks):
#     """Generate a new resume based on template and skill matrix data"""
#     # Validate that files exist
#     if not os.path.exists(request.old_resume_path):
#         raise HTTPException(status_code=400, detail="Resume file not found. Please upload a resume first.")
    
#     if not os.path.exists(request.template_path):
#         raise HTTPException(status_code=400, detail="Template file not found. Please upload a template first.")
    
#     try:
#         # Extract text from files
#         old_resume_text = extract_text_from_file(request.old_resume_path)
#         template_text = extract_text_from_file(request.template_path)
        
#         # Get skill matrix data if available
#         skill_matrix_json = "No skill matrix data available"
#         if request.skill_matrix_path and os.path.exists(request.skill_matrix_path):
#             # If first name and last name are provided, get specific data
#             if request.first_name and request.last_name:
#                 skill_matrix_data = get_json_by_name(request.first_name, request.last_name)
#                 if skill_matrix_data:
#                     skill_matrix_json = json.dumps(skill_matrix_data)
#             else:
#                 # Log a warning but don't fail
#                 print("Warning: Skill matrix provided but no name specified")
        
#         # Step 1: Extract structured JSON from template
#         try:
#             template_json_str = LLM_CALL_1(template_text)
            
#             # Clean up the JSON string to ensure it's valid
#             template_json_str = template_json_str.replace('```json', '').replace('```', '').strip()
            
#             try:
#                 formatted_json_schema = json.loads(template_json_str)
#             except json.JSONDecodeError as e:
#                 # If template JSON parsing fails, use a default schema
#                 print(f"Error parsing template JSON: {e}, using default schema")
#                 formatted_json_schema = {
#                     "name": "",
#                     "designation": "",
#                     "objective": "",
#                     "education": [],
#                     "skills": [],
#                     "project_details": {}
#                 }
#         except Exception as e:
#             print(f"Error in LLM_CALL_1: {str(e)}")
#             # Use default schema if LLM call fails
#             formatted_json_schema = {
#                 "name": "",
#                 "designation": "",
#                 "objective": "",
#                 "education": [],
#                 "skills": [],
#                 "project_details": {}
#             }
        
#         # Step 2: Reformat old resume with skill matrix data
#         try:
#             new_resume_json = LLM_CALL_2(
#                 formatted_json_schema, 
#                 old_resume_text, 
#                 skill_matrix_json
#             )
            
#             # Ensure we have valid JSON
#             try:
#                 resume_data = json.loads(new_resume_json)
#             except json.JSONDecodeError as e:
#                 # Clean up the JSON if necessary
#                 print(f"Error parsing generated JSON: {e}, attempting to clean")
#                 clean_json = new_resume_json.replace('```json', '').replace('```', '').strip()
#                 try:
#                     resume_data = json.loads(clean_json)
#                 except json.JSONDecodeError:
#                     # If still failing, create minimal valid JSON
#                     print("Error parsing cleaned JSON, using minimal data")
#                     # Extract name from resume
#                     name = "Candidate"
#                     for line in old_resume_text.split('\n'):
#                         if line.strip() and len(line.strip()) < 50:  # Simple heuristic to find a name
#                             name = line.strip()
#                             break
                    
#                     resume_data = {
#                         "name": name,
#                         "designation": "Professional",
#                         "objective": "Experienced professional seeking new opportunities.",
#                         "education": ["Education details not extracted"],
#                         "skills": ["Skills not extracted"],
#                         "project_details": {
#                             "project1": {
#                                 "name": "Project",
#                                 "role": "Team Member",
#                                 "description": "Project description not extracted",
#                                 "technology": "Technologies not extracted",
#                                 "role_played": "Role details not extracted"
#                             }
#                         }
#                     }
#         except Exception as e:
#             print(f"Error in LLM_CALL_2: {str(e)}")
#             # Create minimal valid JSON if LLM call fails
#             name = "Candidate"
#             for line in old_resume_text.split('\n'):
#                 if line.strip() and len(line.strip()) < 50:
#                     name = line.strip()
#                     break
                
#             resume_data = {
#                 "name": name,
#                 "designation": "Professional",
#                 "objective": "Experienced professional seeking new opportunities.",
#                 "education": ["Education details not extracted"],
#                 "skills": ["Skills not extracted"],
#                 "project_details": {
#                     "project1": {
#                         "name": "Project",
#                         "role": "Team Member",
#                         "description": "Project description not extracted",
#                         "technology": "Technologies not extracted",
#                         "role_played": "Role details not extracted"
#                     }
#                 }
#             }
            
#         # Step 3: Generate the resume PDF
#         try:
#             resume_path = generate_resume_1(resume_data)
#         except Exception as e:
#             print(f"Error generating PDF: {str(e)}")
#             # Create a simple PDF with basic information
#             pdf = PDF()
#             pdf.set_margins(15, 15, 15)
#             pdf.add_page()
#             pdf.set_font("Arial", "B", 16)
#             pdf.cell(0, 8, resume_data.get("name", "Candidate"), ln=True, align="L")
#             pdf.set_font("Arial", "", 11)
#             pdf.multi_cell(0, 7, "An error occurred while generating the complete resume.")
#             pdf.ln(10)
#             pdf.multi_cell(0, 7, "Please try again or contact support if the issue persists.")
            
#             resume_path = os.path.join(OUTPUT_DIR, f"{resume_data.get('name', 'Candidate').replace(' ', '_')}_Resume.pdf")
#             pdf.output(resume_path)
        
#         # Step 4: Generate a cover letter
#         cover_letter_status = "Not generated"
#         try:
#             cover_letter_text = generate_cover_letter_from_resume(resume_data)
#             cover_letter_path = generate_cover_letter_pdf(cover_letter_text, resume_data["name"])
#             cover_letter_status = f"Generated successfully"
#         except Exception as e:
#             cover_letter_status = f"Failed to generate"
#             print(f"Error generating cover letter: {str(e)}")
        
#         return {
#             "message": "Resume generated successfully",
#             "resume_path": resume_path,
#             "cover_letter_status": cover_letter_status
#         }
    
#     except Exception as e:
#         import traceback
#         error_detail = str(e)
#         print(f"Error in generate_resume endpoint: {error_detail}")
#         print(traceback.format_exc())
#         raise HTTPException(status_code=500, detail=f"Error generating resume: {error_detail}")

    




@app.post("/generate/resume", response_model=dict)
async def generate_resume(request: ResumeRequest, background_tasks: BackgroundTasks):
    """Generate a new resume based on template and input files"""
    # Validate that template exists
    if not os.path.exists(request.template_path):
        raise HTTPException(status_code=400, detail="Template file not found. Please upload a template first.")
    
    try:
        # Extract template text
        template_text = extract_text_from_file(request.template_path)
        
        # Step 1: Extract structured JSON from template
        try:
            template_json_str = LLM_CALL_1(template_text)
            
            # Clean up the JSON string to ensure it's valid
            template_json_str = template_json_str.replace('```json', '').replace('```', '').strip()
            
            try:
                formatted_json_schema = json.loads(template_json_str)
            except json.JSONDecodeError as e:
                # If template JSON parsing fails, use a default schema
                print(f"Error parsing template JSON: {e}, using default schema")
                formatted_json_schema = {
                    "name": "",
                    "designation": "",
                    "objective": "",
                    "education": [],
                    "skills": [],
                    "project_details": {}
                }
        except Exception as e:
            print(f"Error in LLM_CALL_1: {str(e)}")
            # Use default schema if LLM call fails
            formatted_json_schema = {
                "name": "",
                "designation": "",
                "objective": "",
                "education": [],
                "skills": [],
                "project_details": {}
            }
        
        # Get input files content if available
        old_resume_text = None
        old_cover_letter_text = None
        skill_matrix_json = None
        
        if request.old_resume_path and os.path.exists(request.old_resume_path):
            old_resume_text = extract_text_from_file(request.old_resume_path)
        
        if request.old_cover_letter_path and os.path.exists(request.old_cover_letter_path):
            old_cover_letter_text = extract_text_from_file(request.old_cover_letter_path)
        
        if request.skill_matrix_path and os.path.exists(request.skill_matrix_path):
            # If first name and last name are provided, get specific data
            if request.first_name and request.last_name:
                skill_matrix_data = get_json_by_name(request.first_name, request.last_name)
                if skill_matrix_data:
                    skill_matrix_json = json.dumps(skill_matrix_data)
            else:
                # Use the whole file content with explicit UTF-8 encoding
                with open(request.skill_matrix_path, 'r', encoding='utf-8', errors='replace') as f:
                    skill_matrix_json = f.read()
                # Use the whole file content
                # with open(request.skill_matrix_path, 'r') as f:
                #     skill_matrix_json = f.read()
        
        # Step 2: Call appropriate LLM function based on available inputs
        try:
            if old_resume_text and old_cover_letter_text and skill_matrix_json:
                # All three inputs available
                resume_content = LLM_CALL_8(formatted_json_schema, old_resume_text, old_cover_letter_text, skill_matrix_json)
            elif old_resume_text and skill_matrix_json:
                # Old resume and skill matrix
                resume_content = LLM_CALL_2(formatted_json_schema, old_resume_text, skill_matrix_json)
            elif old_cover_letter_text and skill_matrix_json:
                # Cover letter and skill matrix
                resume_content = LLM_CALL_3(formatted_json_schema, old_cover_letter_text, skill_matrix_json)
            elif old_resume_text and old_cover_letter_text:
                # Old resume and cover letter
                resume_content = LLM_CALL_4(formatted_json_schema, old_resume_text, old_cover_letter_text)
            elif old_resume_text:
                # Only old resume
                resume_content = LLM_CALL_5(formatted_json_schema, old_resume_text)
            elif old_cover_letter_text:
                # Only cover letter
                resume_content = LLM_CALL_6(formatted_json_schema, old_cover_letter_text)
            elif skill_matrix_json:
                # Only skill matrix
                resume_content = LLM_CALL_7(formatted_json_schema, skill_matrix_json)
            else:
                raise HTTPException(status_code=400, detail="At least one additional file (skill matrix, old resume, or old cover letter) is required.")
                
            # Ensure we have valid JSON
            try:
                resume_data = json.loads(resume_content)
            except json.JSONDecodeError as e:
                # Clean up the JSON if necessary
                print(f"Error parsing generated JSON: {e}, attempting to clean")
                clean_json = resume_content.replace('```json', '').replace('```', '').strip()
                try:
                    resume_data = json.loads(clean_json)
                except json.JSONDecodeError:
                    # If still failing, create minimal valid JSON
                    print("Error parsing cleaned JSON, using minimal data")
                    # Extract name from resume if available
                    name = "Candidate"
                    if old_resume_text:
                        for line in old_resume_text.split('\n'):
                            if line.strip() and len(line.strip()) < 50:  # Simple heuristic to find a name
                                name = line.strip()
                                break
                    
                    resume_data = {
                        "name": name,
                        "designation": "Professional",
                        "objective": "Experienced professional seeking new opportunities.",
                        "education": ["Education details not extracted"],
                        "skills": ["Skills not extracted"],
                        "project_details": {
                            "project1": {
                                "name": "Project",
                                "role": "Team Member",
                                "description": "Project description not extracted",
                                "technology": "Technologies not extracted",
                                "role_played": "Role details not extracted"
                            }
                        }
                    }
        except Exception as e:
            print(f"Error in LLM call: {str(e)}")
            # Create minimal valid JSON if LLM call fails
            name = "Candidate"
            if old_resume_text:
                for line in old_resume_text.split('\n'):
                    if line.strip() and len(line.strip()) < 50:
                        name = line.strip()
                        break
                    
            resume_data = {
                "name": name,
                "designation": "Professional",
                "objective": "Experienced professional seeking new opportunities.",
                "education": ["Education details not extracted"],
                "skills": ["Skills not extracted"],
                "project_details": {
                    "project1": {
                        "name": "Project",
                        "role": "Team Member",
                        "description": "Project description not extracted",
                        "technology": "Technologies not extracted",
                        "role_played": "Role details not extracted"
                    }
                }
            }
            
        # Step 3: Generate the resume PDF
        try:
            resume_path = generate_resume_1(resume_data)
        except Exception as e:
            print(f"Error generating PDF: {str(e)}")
            # Create a simple PDF with basic information
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, resume_data.get("name", "Candidate"), ln=True, align="L")
            pdf.set_font("Arial", "", 11)
            pdf.multi_cell(0, 7, "An error occurred while generating the complete resume.")
            pdf.ln(10)
            pdf.multi_cell(0, 7, "Please try again or contact support if the issue persists.")
            
            resume_path = os.path.join(OUTPUT_DIR, f"{resume_data.get('name', 'Candidate').replace(' ', '_')}_Resume.pdf")
            pdf.output(resume_path)
        
        # Step 4: Generate a cover letter
        cover_letter_status = "Not generated"
        try:
            cover_letter_text = generate_cover_letter_from_resume(resume_data)
            cover_letter_path = generate_cover_letter_pdf(cover_letter_text, resume_data["name"])
            cover_letter_status = "Generated successfully"
        except Exception as e:
            cover_letter_status = f"Failed to generate: {str(e)}"
            print(f"Error generating cover letter: {str(e)}")
        
        return {
            "message": "Resume generated successfully",
            "resume_path": resume_path,
            "cover_letter_status": cover_letter_status
        }
    
    except Exception as e:
        import traceback
        error_detail = str(e)
        print(f"Error in generate_resume endpoint: {error_detail}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Error generating resume: {error_detail}")


@app.get("/download/{filename}")
async def download_file(filename: str):
    """Endpoint to download generated files"""
    file_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(file_path, filename=filename)

# Create prompts if they don't exist
def create_prompts():
    llm1_path = os.path.join(PROMPTS_DIR, "LLM1.txt")
    if not os.path.exists(llm1_path):
        with open(llm1_path, "w") as f:
            f.write("""
You are an AI assistant that extracts structured information from resumes.
Your task is to analyze the provided resume text and extract key details into a well-structured JSON format.

The JSON structure should include:
{
  "name": "Full Name",
  "designation": "Job Title/Role",
  "objective": "Professional summary/objective statement",
  "education": ["Degree 1", "Degree 2"],
  "skills": ["Skill 1", "Skill 2", "Skill 3"],
  "project_details": {
    "project1": {
      "name": "Project Name",
      "role": "Role in Project",
      "description": "Brief description of the project",
      "technology": "Technologies used",
      "role_played": "Responsibilities and contributions"
    },
    "project2": {
      "name": "Project Name",
      "role": "Role in Project",
      "description": "Brief description of the project",
      "technology": "Technologies used",
      "role_played": "Responsibilities and contributions"
    }
  }
}

Follow these guidelines:
1. Extract information ONLY from the provided text.
2. Do not invent or add details that aren't present in the original text.
3. Format the output as valid JSON that strictly adheres to the structure above.
4. If certain information is missing, use empty strings or arrays rather than omitting fields.
5. Include at least 2-3 projects if available in the resume.
6. Organize skills in a logical manner, grouping similar skills together.
7. Be concise but comprehensive in extracting information.
8. Ensure all JSON is properly formatted and valid.
            """)

# Run startup tasks
@app.on_event("startup")
async def startup_event():
    create_prompts()

# Mount static files directory
# app.mount("/static", StaticFiles(directory=os.path.join(BASE_PATH, "static")), name="static")
static_dir = os.path.join(BASE_PATH, "static")
os.makedirs(static_dir, exist_ok=True)  # Create the directory if it doesn't exist
app.mount("/static", StaticFiles(directory=static_dir), name="static")

# Run the application
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

